#!/usr/bin/env python3
"""Genera el informe breve de validación y confiabilidad en formato Word.

Produce un .docx con carátula institucional, semáforo de color por criterio,
porcentajes destacados y las figuras del análisis incrustadas. El contenido es
el mismo de `docs/entregables/02-validacion-y-confiabilidad/`; aquí se
le da la presentación formal que pide el curso.

Uso:
    .venv/bin/python3 scripts/entregables/generar_informe_word.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[2]
GRAF = REPO / "docs" / "entregables" / "graficas"
OUT = (REPO / "docs" / "entregables" / "02-validacion-y-confiabilidad"
       / "Informe-validacion-confiabilidad.docx")
LOGO = REPO / "docs" / "entregables" / "assets" / "logo-upeu.png"

# Paleta coherente con las gráficas del informe
INK = RGBColor(0x13, 0x1B, 0x2E)
DIM = RGBColor(0x5B, 0x6B, 0x8C)
ACCENT = RGBColor(0x0F, 0x8A, 0x7D)
OK = RGBColor(0x15, 0x80, 0x3D)
AMBER = RGBColor(0xB4, 0x53, 0x09)
DANGER = RGBColor(0xB9, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Rellenos de celda (hex sin #)
F_OK, F_AMBER, F_DANGER = "E0F3E6", "FDECD2", "FBE3E1"
F_HEAD, F_ZEBRA = "0F8A7D", "F4F6FB"


# --------------------------------------------------------------- utilidades --
def shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def cell_text(cell, text, *, bold=False, color=INK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def tabla(doc, headers, rows, widths=None, header_fill=F_HEAD):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_text(c, h, bold=True, color=WHITE, size=9.5)
        shade(c, header_fill)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            fill = None
            if isinstance(val, tuple):
                val, fill = val
            cell_text(cells[i], str(val))
            if fill:
                shade(cells[i], fill)
            elif r_i % 2 == 1:
                shade(cells[i], F_ZEBRA)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


def h1(doc, texto, numero=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    if numero:
        r = p.add_run(f"{numero}  ")
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = ACCENT
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = INK
    # línea inferior
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "0F8A7D")
    b.append(bottom)
    pPr.append(b)
    return p


def h2(doc, texto, color=INK, icono=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{icono}{texto}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = color
    return p


def parrafo(doc, texto, *, size=10, color=INK, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    partes = texto.split("**")
    for i, seg in enumerate(partes):
        if not seg:
            continue
        r = p.add_run(seg)
        r.bold = i % 2 == 1
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p


def vineta(doc, texto, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    partes = texto.split("**")
    for i, seg in enumerate(partes):
        if not seg:
            continue
        r = p.add_run(seg)
        r.bold = i % 2 == 1
        r.font.size = Pt(10)
        r.font.color.rgb = color
    return p


def figura(doc, nombre, pie, ancho=15.5):
    doc.add_picture(str(GRAF / nombre), width=Cm(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(pie)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = DIM


def caja(doc, titulo, texto, color_borde="0F8A7D", fill="EEF1F8"):
    """Recuadro destacado, implementado como tabla de una celda."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(color_borde)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i, seg in enumerate(texto.split("**")):
        if not seg:
            continue
        rr = p2.add_run(seg)
        rr.bold = i % 2 == 1
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = INK
    return t


# ------------------------------------------------------------------ informe --
def main() -> int:
    man = json.loads((REPO / "artifacts/model/manifest.json").read_text())
    o = man["evaluation"]["ocsvm_scaled"]
    fpr = o["test"]["fpr"] * 100
    det = o["anomalies"]["detection_rate"] * 100
    kali = o["anomalies"]["kali_real_detection_rate"] * 100

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)

    # ---------------------------------------------------------- CARÁTULA ----
    if not LOGO.exists():
        raise SystemExit(f"falta el logo: {LOGO}")
    doc.add_picture(str(LOGO), width=Cm(7.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sz, col, bold in [
        ("Universidad Peruana Unión", 13, INK, True),
        ("Facultad de Ingeniería y Arquitectura", 11, DIM, False),
        ("E.P. de Ingeniería de Sistemas", 11, DIM, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = col

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME DE VALIDACIÓN INTERNA, VALIDACIÓN EXTERNA\nY CONFIABILIDAD DE LOS RESULTADOS")
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Detección temprana de comportamientos anómalos en redes de datos\n"
                  "mediante modelos predictivos y un mecanismo de control inline")
    r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = DIM

    doc.add_paragraph()
    tabla(doc,
          ["Campo", "Detalle"],
          [("Curso", "Investigación V · Ciclo X"),
           ("Docente", "Ing. Nemias Saboya Ríos"),
           ("Integrantes", "Rubén Mark Salazar Tocas\nUziel Elias Sauñe Fernandez"),
           ("Asesores", "Ing. Nemias Saboya Ríos · Ing. Fernando Manuel Asin Gómez"),
           ("Fecha", "19 de agosto de 2026")],
          widths=[4.0, 12.5])

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Lima, agosto de 2026"); r.font.size = Pt(10.5); r.font.color.rgb = DIM
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -------------------------------------------------- 1. ESTADO GENERAL ---
    h1(doc, "Estado general de los resultados", "1.")
    parrafo(doc, "Se evaluó de forma rigurosa si los resultados obtenidos son **válidos y confiables**, "
                 "aplicando los tres criterios solicitados. Cada criterio indica qué se abordó de manera "
                 "concreta, qué parcialmente y qué no se abordó.")

    tabla(doc,
          ["Criterio", "Estado", "Síntesis"],
          [("Validación interna", ("PARCIAL", F_AMBER),
            "Controles anti-fuga reales y verificados, pero el modelo final se eligió observando el conjunto de prueba"),
           ("Validación externa", ("INSUFICIENTE", F_DANGER),
            "Medida en operación real y refutada: el error sobre tráfico legítimo intenso es 5 veces el de laboratorio"),
           ("Confiabilidad", ("ALTA", F_OK),
            "Resultados reproducibles: al repetir la evaluación se obtienen exactamente las mismas cifras")],
          widths=[3.8, 2.7, 10.0])

    doc.add_paragraph()
    caja(doc, "Lectura general",
         f"Los resultados **sí sostienen** que el sistema detecta y bloquea ataques reales en tiempo real "
         f"(detección del **{kali:.1f} %** sobre ataques genuinos, ROC-AUC de **0,974**, bloqueo en una mediana "
         f"de **8 segundos**). **No sostienen todavía** que lo haga sin penalizar el tráfico legítimo intenso. "
         f"La debilidad principal no está en la ingeniería del sistema, sino en el diseño estadístico de la evaluación.")

    # --------------------------------------------- 2. VALIDACIÓN INTERNA ---
    doc.add_paragraph()
    h1(doc, "Validación interna", "2.")
    parrafo(doc, "¿Los resultados obtenidos se deben realmente a lo que el estudio dice haber probado?",
            italic=True, color=DIM)

    h2(doc, "Abordado de manera concreta", OK, "✔  ")
    tabla(doc,
          ["Aspecto", "Evidencia verificable"],
          [("Sin fuga de datos entre particiones",
            "Auditoría automática: ningún episodio se reparte entre entrenamiento, validación y prueba"),
           ("Sin información futura en las variables",
            "Prueba unitaria: un evento posterior no altera una ventana ya cerrada"),
           ("Umbral fijado antes de ver la prueba",
            "Cuantil α = 0,05 solo sobre validación (k = 13, n = 273); el escalador se ajusta solo con entrenamiento"),
           ("Evaluación en un solo paso",
            "Sin reentrenar tras ver resultados; registro sellado con hash del calibrador y del repositorio"),
           ("Detección de una fuga propia",
            "Un experimento con selección contaminada fue identificado y marcado como “no debe citarse”")],
          widths=[5.2, 11.3])

    h2(doc, "Abordado parcialmente", AMBER, "◐  ")
    vineta(doc, "**Análisis de sensibilidad.** Se ejecutó con 10 semillas, ponderación por episodio y "
                "colapso de duplicados, pero **solo sobre los modelos descartados**. El modelo finalmente "
                "elegido no recibió ninguna prueba de estabilidad.")

    h2(doc, "No abordado", DANGER, "✘  ")
    vineta(doc, "**Selección del modelo sin contaminar la prueba.** El modelo se eligió después de observar "
                "su desempeño en el conjunto de prueba. El registro del proyecto documenta que ese modelo "
                "estaba designado como “comparador” y que la política prohibía promoverlo por ganar una "
                f"métrica posterior. En consecuencia, el **{det:.1f} %** de detección es el máximo entre "
                "**7 candidatos** evaluados sobre los mismos datos, sin conjunto reservado que permita una "
                "estimación sin sesgo optimista.")
    vineta(doc, "**Pruebas de significancia estadística.** No se realizó ninguna prueba (t, Wilcoxon o "
                "equivalente) que compare los modelos entre sí.")

    # --------------------------------------------- 3. VALIDACIÓN EXTERNA ---
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h1(doc, "Validación externa", "3.")
    parrafo(doc, "¿Los resultados se generalizan a otros contextos, poblaciones o condiciones de uso?",
            italic=True, color=DIM)

    h2(doc, "Abordado de manera concreta", OK, "✔  ")
    vineta(doc, "**Se midió el sistema completo en operación real**, no solo el modelo en laboratorio: "
                "2 pases de 29 corridas más 2 pruebas de aislamiento, con motor y bloqueo activos.")
    vineta(doc, "**Ataques genuinos** desde una máquina atacante real en 6 familias distintas, no simulados "
                "por inyección de datos.")
    vineta(doc, "**Se declaró la procedencia heterogénea** de los datos de ataque: 161 ventanas reales y "
                f"18 heredadas, reportadas por separado (**{kali:.1f} %** frente a **83,3 %**).")

    h2(doc, "Resultado que refuta la generalización", DANGER, "✘  ")
    parrafo(doc, "Es el hallazgo más importante del informe y se reporta aunque sea desfavorable:")
    tabla(doc,
          ["Condición de medición", "Falsos positivos", "IC 95 %"],
          [("Laboratorio (conjunto de prueba)", (f"{fpr:.2f} %  (13/276)", F_OK), "2,8 % – 7,9 %"),
           ("Operación real · pase 1", ("25,8 %  (16/62)", F_DANGER), "16,6 % – 37,9 %"),
           ("Operación real · pase 2", ("23,0 %  (17/74)", F_DANGER), "14,9 % – 33,7 %")],
          widths=[6.5, 5.0, 5.0])
    doc.add_paragraph()
    figura(doc, "C1-fpr-offline-vs-operativo.png",
           "Figura 1. Los intervalos de confianza no se solapan: la diferencia no se explica por azar muestral.")

    caja(doc, "Evidencia adicional en aislamiento",
         "Se reprodujo **sin contaminación entre pruebas**: una transferencia legítima de 200 Mbit/s generó una "
         "ventana que cruzó el umbral y **bloqueó a un cliente legítimo durante 120 segundos**. Otra ventana de "
         "la misma transferencia se permitió por apenas **0,0014 puntos** de score, lo que indica que el tráfico "
         "legítimo intenso cae dentro del margen de decisión del modelo.",
         color_borde="B91C1C", fill="FBE3E1")

    doc.add_paragraph()
    h2(doc, "No abordado", DANGER, "✘  ")
    vineta(doc, "**Partición por sesiones independientes.** La división se hizo por índice de repetición, "
                "por lo que los **44 perfiles** de tráfico aparecen en las tres particiones: se mide "
                "repetibilidad del escenario, no generalización a tráfico no visto.")
    vineta(doc, "**Jornada de validación temporal externa.** No existe un conjunto capturado en fecha distinta "
                "y reservado sin participar en entrenamiento ni calibración.")
    vineta(doc, "**Diversidad de escenarios.** Faltan seis escenarios legítimos previstos (SSH, SCP/SFTP, SMB, "
                "respaldo, streaming y actualizaciones) y no hay captura multi-sistema-operativo.")

    # ------------------------------------------------- 4. CONFIABILIDAD ----
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h1(doc, "Confiabilidad", "4.")
    parrafo(doc, "¿Repetir el procedimiento produce los mismos resultados?", italic=True, color=DIM)

    h2(doc, "Abordado de manera concreta", OK, "✔  ")
    tabla(doc,
          ["Aspecto", "Evidencia"],
          [("Reproducibilidad verificada",
            "Al reevaluar el modelo congelado se obtuvieron exactamente las mismas cifras del registro "
            "original (13/276 y 158/179)"),
           ("Integridad de artefactos",
            "SHA-256 de datos, modelo y programa de calibración; repositorio verificado limpio antes y después"),
           ("Trazabilidad",
            "330 registros de cambios · 181 documentos de campañas · 162 revisiones independientes"),
           ("Estabilidad operativa",
            "100 % de disponibilidad en 57 corridas, sin pérdida de paquetes en captura"),
           ("Consistencia entre repeticiones",
            "Los dos pases de validación operativa dieron resultados equivalentes (25,8 % y 23,0 %)")],
          widths=[5.2, 11.3])

    h2(doc, "Abordado parcialmente", AMBER, "◐  ")
    parrafo(doc, "**Cuantificación de la incertidumbre.** El trabajo original no calculó ninguna medida de "
                 "dispersión. Se incorporan en este informe **intervalos de confianza de Wilson al 95 %**, que "
                 "revelan un problema que las cifras puntuales ocultaban:")
    tabla(doc,
          ["Cifra reportada", "IC 95 % real", "Lectura"],
          [("50 % de detección en fuerza bruta (3/6)", ("18,8 % – 81,2 %", F_DANGER),
            "Con n = 6 no sostiene ninguna conclusión"),
           ("55,2 % en password-spray (16/29)", ("37,5 % – 71,6 %", F_AMBER),
            "Intervalo muy amplio; conclusión débil"),
           (f"{det:.1f} % de detección global (158/179)", ("82,7 % – 92,2 %", F_OK), "Sólido")],
          widths=[6.5, 4.5, 5.5])

    h2(doc, "No abordado", DANGER, "✘  ")
    vineta(doc, "**Repetición del experimento de modelado.** La calibración se ejecutó una sola vez; no hay "
                "repeticiones independientes que permitan estimar la variabilidad del umbral.")
    vineta(doc, "**Confiabilidad inter-evaluador.** No aplica al diseño actual, que no emplea jueces ni "
                "instrumentos de percepción.")

    # --------------------------------- 5. QUÉ FALTA Y CÓMO SE ABORDARÁ -----
    doc.add_paragraph()
    h1(doc, "Qué falta y cómo se abordará con el tiempo disponible", "5.")
    parrafo(doc, "Ordenado por relación entre costo y beneficio. **Ninguna acción de los bloques A y B "
                 "requiere capturar datos nuevos.**")
    tabla(doc,
          ["#", "Acción", "Corrige", "Tiempo"],
          [(("A1", F_OK), "Declarar la selección posterior del modelo y que la detección reportada es una estimación optimista", "Validez interna", ("Horas", F_OK)),
           (("A2", F_OK), "Incorporar los intervalos de confianza a todas las proporciones reportadas", "Confiabilidad", ("Horas", F_OK)),
           (("A3", F_OK), "Sustituir las conclusiones sobre familias con n ≤ 6 por declaración de muestra insuficiente", "Confiabilidad", ("Horas", F_OK)),
           (("A4", F_OK), "Reportar el error operativo (23–26 %) junto al de laboratorio (4,71 %)", "Validez externa", ("Horas", F_OK)),
           (("A5", F_OK), "Publicar el diccionario de fórmulas de las 14 variables nuevas", "Constructo", ("Horas", F_OK)),
           (("B1", F_AMBER), "Ejecutar la prueba de ablación por capas (L3/L4/L7) y la comparación 14 vs. 28 variables", "Constructo", ("1–2 días", F_AMBER)),
           (("B2", F_AMBER), "Prueba de estabilidad por remuestreo del modelo elegido", "Validez interna", ("Horas", F_AMBER)),
           (("B3", F_AMBER), "Prueba de significancia entre modelos", "Validez interna", ("Horas", F_AMBER)),
           (("C1", F_DANGER), "Capturar una jornada nueva y reservarla como validación temporal externa", "Validez externa", ("Días", F_DANGER)),
           (("C2", F_DANGER), "Recalibrar el umbral incluyendo tráfico legítimo intenso y repetir la validación", "Validez externa", ("1–2 semanas", F_DANGER))],
          widths=[1.2, 9.0, 3.2, 2.6])

    doc.add_paragraph()
    caja(doc, "Compromiso realista",
         "Con el tiempo disponible se ejecutarán los **bloques A y B** antes de cerrar la tesis: cubren los dos "
         "requisitos formales pendientes y corrigen la principal deficiencia estadística **sin requerir "
         "experimentación nueva**. El **bloque C** se declarará como trabajo futuro, indicando con precisión "
         "qué quedaría por demostrar.")

    # ------------------------------------------------------ 6. CONCLUSIÓN --
    doc.add_paragraph()
    h1(doc, "Conclusión", "6.")
    parrafo(doc, f"Los resultados sostienen una afirmación **acotada y verdadera**: se demostró la viabilidad de "
                 f"detectar comportamientos anómalos y ejercer control en línea en tiempo real sobre una red real, "
                 f"con capacidad discriminante alta (**ROC-AUC = 0,974**), detección del **{kali:.1f} %** sobre "
                 f"ataques genuinos y bloqueo en una mediana de **8 segundos**.")
    parrafo(doc, "No sostienen todavía que el sistema sea apto para operación desatendida: sobre tráfico legítimo "
                 "de alto volumen el error alcanza **23–26 %**. Esa limitación **está medida, cuantificada y "
                 "declarada**, que es la condición que la hace defendible ante una revisión por pares.")
    parrafo(doc, "La prioridad antes de cerrar no es mejorar el sistema, sino **corregir la inferencia**: declarar "
                 "la selección posterior del modelo, acompañar cada cifra de su intervalo de confianza y ejecutar "
                 "la ablación pendiente.")

    # ----------------------------------------------------- REFERENCIAS ----
    doc.add_paragraph()
    h1(doc, "Referencias")
    for ref in [
        "**Campbell, D. T., & Stanley, J. C. (1963).** Experimental and quasi-experimental designs for research. "
        "Rand McNally. — Marco de validez interna y externa aplicado en las secciones 2 y 3.",
        "**Wilson, E. B. (1927).** Probable inference, the law of succession, and statistical inference. Journal of "
        "the American Statistical Association, 22(158), 209–212. — Método de los intervalos de confianza.",
        "**Cronbach, L. J. (1951).** Coefficient alpha and the internal structure of tests. Psychometrika, 16(3), "
        "297–334. — Criterio de consistencia interna, no aplicable por no emplear instrumentos psicométricos.",
        "**Peng, R. D. (2011).** Reproducible research in computational science. Science, 334(6060), 1226–1227. — "
        "Distinción entre reproducibilidad y replicabilidad, sección 4.",
        "**Kapoor, S., & Narayanan, A. (2023).** Leakage and the reproducibility crisis in machine-learning-based "
        "science. Patterns, 4(9), 100804. — Tipología de fugas y del sesgo por selección sobre el conjunto de prueba.",
        "**ISO/IEC 25010:2011.** Systems and software engineering — SQuaRE — System and software quality models. — "
        "Su correspondencia detallada se desarrolla en la ficha de auditoría del producto.",
    ]:
        vineta(doc, ref)

    doc.add_paragraph()
    caja(doc, "Anexo técnico",
         "El análisis completo, con las 11 figuras, la comparación de los 7 modelos evaluados y el detalle de cada "
         "hallazgo, está disponible en el repositorio del proyecto: "
         "docs/entregables/01-evaluacion-critica/",
         fill="EEF1F8")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Generado: {OUT.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
