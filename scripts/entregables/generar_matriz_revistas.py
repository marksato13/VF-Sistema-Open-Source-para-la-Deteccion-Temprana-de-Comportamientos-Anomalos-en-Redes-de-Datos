#!/usr/bin/env python3
"""Matriz de decision de revistas (Sesion 04).

Cada celda lleva su fuente y su estado de verificacion. Los puntajes se
calculan; no se escriben a mano. Un dato sin fuente primaria se marca como
pendiente en vez de presentarse como verificado.

Formula del complemento:  aporte = puntaje x peso / 10   (total sobre 100)

    python3 scripts/entregables/generar_matriz_revistas.py
"""
from __future__ import annotations
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
OUT_MD = REPO / "docs/entregables/09-matriz-revistas/matriz-decision-revistas.md"
FECHA = "26 de agosto de 2026"

# (clave, nombre, peso, regla de puntuacion)
# Los cinco criterios y sus pesos son los de la Sesión 04 (diapositiva 15). El
# nombre local va primero y entre paréntesis el nombre de la sesión, para que la
# equivalencia sea verificable de un vistazo.
CRITERIOS = [
    ("pertinencia", "Pertinencia temática (alcance)", 30,
     "Coincidencia entre el alcance editorial declarado y el problema, método y "
     "dominio del artículo. Puntuado con la evidencia de los 21 artículos semilla"),
    ("visibilidad", "Visibilidad bibliométrica (cuartil)", 25,
     "CiteScore, percentil, cuartil SJR y h5-index de Google Scholar Metrics, "
     "identificados por separado y con su fuente"),
    ("viabilidad", "Viabilidad editorial (tiempo)", 20,
     "Tipo de revisión, tiempo real medido, periodicidad y capacidad de publicación"),
    ("costo", "Costo y accesibilidad (APC)", 15,
     "APC vigente, quién lo cubre, cargos por página y acceso abierto"),
    ("formato", "Compatibilidad formal (requisitos)", 10,
     "Plantilla, extensión exigida y requisitos de envío"),
]

# h5-index y h5-mediana de Google Scholar Metrics, consultados el 02/09/2026.
# La Sesión 04 (diapositiva 16) lo lista como una de las cuatro métricas de
# prestigio. None = la revista no figura en Google Scholar Metrics.
H5 = {
    "IJACSA": (63, 84), "IJIT": (58, 94), "BEEI": (41, 53),
    "ISJ": (27, 38), "ISI": (23, 32), "CIT": (20, 29),
    "IJIES": None,   # ausencia comprobada con tres consultas distintas
    "IJSSE": None,   # ausencia comprobada con dos consultas distintas
    "ICS": None,     # la coincidencia es otra revista: Int. J. of Inf. and Computer Security
}

# ✔ = verificado en fuente primaria · ~ = fuente secundaria · ? = pendiente
CANDIDATAS = [
    {
        "nombre": "Bulletin of Electrical Engineering and Informatics (BEEI)",
        "corto": "BEEI", "editor": "IAES · Indonesia", "issn": "2089-3191",
        "scopus": "https://www.scopus.com/sourceid/21100826382",
        "control": ("En la lista", "N.º 8 de la lista de control institucional, registrada como Q1 · Indonesia · USD 385"),
        "filtro": ("Supera", "ISSN confirmado, ficha activa en Scopus, política de revisión por pares "
                             "publicada, archivo de números con DOI, APC transparente y editor "
                             "identificable. Fuera de la lista de revistas depredadoras consultada"),
        "datos": {
            "CiteScore 2025": ("4,2 · percentil 65 en Computer Networks and Communications", "✔"),
            "Cuartil SJR": ("Q3 · la revista declara además Q1 por CiteScore", "~"),
            "Revisión": ("Single-blind, ≥2 revisores; 8–12 semanas declaradas", "✔"),
            "Periodicidad": ("Bimestral · 76 artículos en el número de agosto de 2026 · "
                             "**377 artículos en 2025**, 453 en 2024 y 415 en 2023", "✔"),
            "APC": ("USD 415 hasta 8 páginas · USD 50 por página adicional · USD 830 si es autor único", "✔"),
            "Plantilla": ("DOCX oficial disponible", "✔"),
        },
        "puntajes": {
            "pertinencia": (9, "Su alcance declara explícitamente redes de comunicaciones, seguridad de "
                               "redes, aprendizaje automático y ciberseguridad: los cuatro ejes del artículo"),
            "visibilidad": (8, "CiteScore 4,2 y percentil 65, el más alto de las candidatas verificadas"),
            "viabilidad": (9, "Bimestral con 76 artículos por número y ciclo declarado de 8 a 12 "
                              "semanas. Su volumen anual (377) es menor que el de IJIES (556)"),
            "costo": (8, "USD 415 con coautoría —USD 830 si fuera autor único—, por encima de "
                         "los USD 300 de IJIES y por debajo de los USD 850 de ISI e IJSSE"),
            "formato": (7, "Plantilla disponible, pero el límite base de 8 páginas obliga a comprimir o "
                           "a pagar por página adicional"),
        },
    },
    {
        "nombre": "International Journal of Safety and Security Engineering (IJSSE)",
        "corto": "IJSSE", "editor": "IIETA · Canadá", "issn": "2041-904X",
        "scopus": "https://www.scopus.com/sourceid/21100785501",
        "control": ("En la lista", "N.º 7 de la lista de control institucional, registrada como Q3 · Reino Unido · USD 700"),
        "filtro": ("Supera", "ISSN confirmado, ficha activa en Scopus, revisión double-blind declarada, "
                             "archivo de números con DOI y página oficial de APC. Fuera de la lista de "
                             "revistas depredadoras consultada"),
        "datos": {
            "CiteScore 2025": ("2,8 · percentil 60 en Safety, Risk, Reliability and Quality", "✔"),
            "Cuartil SJR": ("Q3", "~"),
            "Revisión": ("Double-blind, ≥2 expertos independientes; ~2 meses", "✔"),
            "Periodicidad": ("12 números regulares al año · 20 artículos en el número de julio de 2026", "✔"),
            "APC": ("USD 850 por artículo aceptado, sin cargo por página", "✔"),
            "Plantilla": ("DOCX oficial disponible", "✔"),
        },
        "puntajes": {
            "pertinencia": (9, "Declara seguridad informática, evaluación de amenazas, ciberseguridad y "
                               "protección de infraestructura crítica; publica de forma habitual detección "
                               "de intrusiones con aprendizaje automático"),
            "visibilidad": (6, "CiteScore 2,8 y percentil 60: por debajo de BEEI en ambos indicadores"),
            "viabilidad": (9, "Doce números al año y revisión double-blind de unos dos meses"),
            "costo": (5, "USD 850, el doble que BEEI"),
            "formato": (8, "Plantilla disponible y sin límite estrecho de páginas declarado"),
        },
    },
    {
        "nombre": "Information Security Journal: A Global Perspective",
        "corto": "ISJ", "editor": "Taylor & Francis · Reino Unido", "issn": "1939-3555",
        "scopus": "https://www.scopus.com/sourceid/19700187807",
        "control": ("Fuera", "No figura entre las 17 revistas de la lista de control institucional"),
        "filtro": ("Supera", "Editorial de trayectoria reconocida, ficha activa en Scopus e indexación "
                             "adicional en ESCI. Fuera de la lista de revistas depredadoras consultada"),
        "datos": {
            "CiteScore 2025": ("pendiente de verificar en la ficha de Scopus", "?"),
            "Cuartil SJR": ("Q2 como mejor cuartil · SJR 0,489 · h-index 33", "~"),
            "Revisión": ("pendiente de verificar en las instrucciones para autores", "?"),
            "Periodicidad": ("6 números al año · **36 · 35 · 38 artículos en 2024 · 2025 · 2026**, "
                             "contados sobre el índice completo de dblp", "~"),
            "APC": ("Híbrida: publicar por la vía de suscripción no exige APC. **No figura en DOAJ**, "
                    "lo que es coherente con que no haya migrado a acceso abierto de pago", "~"),
            "Plantilla": ("pendiente de verificar", "?"),
        },
        "puntajes": {
            "pertinencia": (9, "Su alcance nombra seguridad de redes y control de acceso; publica trabajos "
                               "sobre ataques SSH y denegación de servicio HTTP, las mismas familias del corpus"),
            "visibilidad": (8, "Mejor cuartil Q2 y h-index 33, el más alto tras Emerald"),
            "viabilidad": (6, "**Corregido**: la cifra anterior de 21 artículos al año era falsa. El "
                              "índice de dblp da 35–38, capacidad equivalente a la de CIT, repartida en "
                              "6 números al año en vez de 4. Baja de 10 porque el tipo y el tiempo de "
                              "revisión siguen sin verificar: dos de los cuatro componentes del criterio"),
            "costo": (10, "Sin APC obligatorio por la vía de suscripción. El acceso abierto "
                          "opcional cuesta USD 3.390, pero nadie está obligado a tomarlo"),
            "formato": (7, "Editorial mayor con formato estándar; requisitos concretos sin verificar"),
        },
    },
    {
        "nombre": "Information and Computer Security",
        "corto": "ICS", "editor": "Emerald · Reino Unido", "issn": "2056-4961",
        "scopus": "https://www.scopus.com/sourceid/21100421900",
        "control": ("Fuera", "No figura entre las 17 revistas de la lista de control institucional"),
        "filtro": ("Supera", "Editorial de trayectoria reconocida y ficha activa en Scopus. Fuera de la "
                             "lista de revistas depredadoras consultada"),
        "datos": {
            "CiteScore 2025": ("pendiente de verificar en la ficha de Scopus", "?"),
            "Cuartil SJR": ("Q2 como mejor cuartil · Q3 en Computer Networks and Communications, "
                            "Information Systems y Software · h-index 60", "~"),
            "Revisión": ("pendiente de verificar", "?"),
            "Periodicidad": ("**34 · 38 · 42 artículos en 2023 · 2024 · 2025**, contados sobre el "
                             "índice completo de dblp", "~"),
            "APC": ("Híbrida: no exige APC por la vía de suscripción. **No figura en DOAJ**, lo que "
                    "es coherente con que no haya migrado a acceso abierto de pago", "~"),
            "Plantilla": ("pendiente de verificar", "?"),
        },
        "puntajes": {
            "pertinencia": (6, "Cubre la categoría de redes, pero su centro editorial se inclina a factores "
                               "humanos, concienciación y cumplimiento de políticas: un artículo puramente "
                               "técnico corre riesgo de quedar fuera de foco"),
            "visibilidad": (9, "h-index 60, el más alto de todas las candidatas"),
            "viabilidad": (6, "Capacidad creciente y la mayor de las tres disponibles —34, 38 y 42 "
                              "artículos en 2023, 2024 y 2025—, pero el tipo y el tiempo de revisión "
                              "siguen sin verificar"),
            "costo": (10, "Sin APC obligatorio por la vía de suscripción. El acceso abierto "
                          "opcional cuesta USD 3.390, pero nadie está obligado a tomarlo"),
            "formato": (6, "Requisitos sin verificar"),
        },
    },
    {
        "nombre": "International Journal of Advanced Computer Science and Applications (IJACSA)",
        "corto": "IJACSA", "editor": "TheSAI · Reino Unido", "issn": "2158-107X",
        "scopus": "https://www.scopus.com/sourceid/21100867241",
        "control": ("En la lista", "N.º 5 de la lista de control institucional, registrada como Q3 · Reino Unido · USD 1 100"),
        "filtro": ("Supera con reserva", "Ficha activa en Scopus e indexación en WoS ESCI, y fuera de la "
                                         "lista de depredadoras consultada. Se registra que su reputación "
                                         "editorial es más discutida que la de las demás candidatas"),
        "datos": {
            "CiteScore 2025": ("3,4", "~"),
            "Cuartil SJR": ("Q3", "~"),
            "Revisión": ("Doble ciego con al menos tres revisores; decisión en unas 3 semanas", "~"),
            "Periodicidad": ("Mensual, con fecha de cierre fija cada mes", "~"),
            "APC": ("GBP 800 · GBP 750 para estudiantes y revisores", "~"),
            "Plantilla": ("Plantilla propia obligatoria", "~"),
        },
        "puntajes": {
            "pertinencia": (6, "Declara cubrir «todas las ramas de las ciencias de la computación»: alcance "
                               "amplio y por tanto menos específico que las tres primeras"),
            "visibilidad": (8, "CiteScore 3,4 y doble indexación en Scopus y WoS ESCI"),
            "viabilidad": (10, "Mensual y con decisión declarada en unas tres semanas: la más rápida"),
            "costo": (3, "GBP 800, el más alto de todas las candidatas"),
            "formato": (7, "Plantilla propia obligatoria, que exige reformatear"),
        },
    },
    {
        "nombre": "Cybernetics and Information Technologies (CIT)",
        "corto": "CIT",
        "editor": "Instituto de TIC · Academia Búlgara de Ciencias · Bulgaria",
        "issn": "1314-4081",
        "scopus": "https://www.scopus.com/sourceid/21100199814",
        "control": ("Fuera", "No figura entre las 17 revistas de la lista de control institucional"),
        "filtro": ("Supera", "ISSN confirmado, ficha activa en Scopus, indexación adicional en Web of "
                             "Science ESCI, DOAJ, INSPEC, ACM Digital Library y Engineering Village; "
                             "adhesión declarada a las Core Practices de COPE, revisión por pares en "
                             "tres fases con cribado antiplagio iThenticate y archivo de números desde "
                             "2001. Fuera de la lista de revistas depredadoras consultada"),
        "datos": {
            "CiteScore 2025": ("4,2 · SNIP 0,854 · la revista publica ambos en su página de indexación", "✔"),
            "Cuartil SJR": ("**Q2** · SJR 0,456 · h-index 27 · además Web of Science ESCI con factor "
                            "de impacto 1,7 (Q3)", "✔"),
            "Revisión": ("Single-blind, ≥2 revisores, en tres fases · **3–6 meses** en modo regular · "
                         "1–3 meses en vía rápida · cribado de texto generado por IA con rechazo "
                         "sin derecho a revisión por encima del 20 %", "✔"),
            "Periodicidad": ("4 números al año · **46 · 43 · 42 artículos en 2023 · 2024 · 2025** "
                             "y 23 en los dos primeros números de 2026", "✔"),
            "APC": ("**600 EUR a partir del 1 de septiembre de 2026** · vía rápida 1 200 EUR · "
                    "solo se paga si el artículo se acepta", "✔"),
            "Plantilla": ("Plantilla DOCX oficial · hasta **20 páginas** · envío por correo electrónico", "✔"),
        },
        "puntajes": {
            "pertinencia": (9, "Su alcance declara tecnologías de comunicación entre computadoras, "
                               "aprendizaje profundo y automático, y reconocimiento de patrones. Lo "
                               "decisivo no es el alcance sino lo que publica: entre 2024 y 2026 sacó "
                               "detección de intrusiones con IA explicable, detección de intrusiones con "
                               "redes convolucionales, detección de DDoS entre conjuntos de datos y una "
                               "arquitectura de **detección y mitigación** de DDoS, que es el análogo "
                               "más cercano al control inline de este proyecto entre todas las candidatas"),
            "visibilidad": (9, "Única candidata con cuartil **Q2 verificado en fuente primaria** y, a la "
                               "vez, factor de impacto en Web of Science. CiteScore 4,2, igual al de BEEI, "
                               "pero con SJR 0,456 frente al Q3 de BEEI"),
            "viabilidad": (5, "El punto débil: 3–6 meses hasta la primera decisión, el plazo más largo "
                              "del conjunto. Su capacidad, en cambio, es estable y comparable a la de "
                              "ISJ e ICS: 46, 43 y 42 artículos en 2023, 2024 y 2025, contados sobre "
                              "los PDF de cada número"),
            "costo": (6, "600 EUR desde septiembre de 2026, por debajo de los USD 850 de IJSSE pero por "
                         "encima de los USD 415 de BEEI"),
            "formato": (9, "Límite de 20 páginas, frente a las 8 de BEEI: el artículo cabe entero sin "
                           "pagar por página adicional. Plantilla oficial y envío por correo"),
        },
    },
    {
        "nombre": "International Journal of Intelligent Engineering and Systems (IJIES)",
        "corto": "IJIES", "editor": "INASS · Japón", "issn": "2185-3118",
        "scopus": "https://www.scopus.com/sourceid/21100199790",
        "control": ("Fuera", "No figura entre las 17 revistas de la lista de control institucional"),
        "filtro": ("Supera con reserva", "ISSN confirmado, ficha en Scopus, política de revisión y "
                   "tarifas publicadas, y **publica su tasa de aceptación**, transparencia poco "
                   "habitual. Fuera de la lista de depredadoras. **La reserva es concreta**: hay que "
                   "confirmar que su cobertura en Scopus sigue activa, porque TELKOMNIKA, IJECE e "
                   "IJEECS —revistas del mismo perfil de alto volumen— fueron descontinuadas en 2025"),
        "datos": {
            "CiteScore 2025": ("3,3 · percentil 62 en General Engineering y 52 en General Computer "
                               "Science, **Q2 en ambas**", "✔"),
            "Cuartil SJR": ("pendiente de verificar en Scimago, hoy bloqueado por Cloudflare", "?"),
            "Revisión": ("Publicación unos 2 meses después de la aceptación · **tasa de aceptación "
                         "declarada: 17,9 % en 2025, 17,8 % en 2024, 14,1 % en 2023**", "✔"),
            "Periodicidad": ("**Mensual desde 2025** · **481 · 556 · 542 artículos** en 2024 · 2025 · "
                             "2026, el mayor volumen de las candidatas disponibles", "✔"),
            "APC": ("**USD 300** · **USD 400 desde el 1 de octubre de 2026** · USD 100 extra si no se "
                    "usa su plantilla · USD 50 por página a partir de la 10.ª", "✔"),
            "Plantilla": ("`IJIES_Format.docx` obligatoria · **límite de 10 páginas** antes del "
                          "recargo", "✔"),
        },
        "puntajes": {
            "pertinencia": (9, "Su alcance nombra ingeniería de redes y computación inteligente, pero "
                               "lo que decide es su producción: **86 artículos desde 2024 con "
                               "«intrusion detection», «anomaly detection» o «network security» en el "
                               "título**. Ninguna otra candidata se acerca a ese volumen temático"),
            "visibilidad": (8, "CiteScore 3,3 y Q2 en sus dos categorías. Por debajo del 4,2 de BEEI y "
                               "CIT, pero con cuartil verificado en la propia revista"),
            "viabilidad": (10, "Mensual, con el mayor volumen de las candidatas de acceso abierto "
                               "—556 al año, frente a 377 de BEEI— y publicación unos 2 meses "
                               "después de aceptar. **Y publica su tasa de aceptación**, que ninguna "
                               "otra hace: 17,9 %"),
            "costo": (9, "USD 300 hoy y USD 400 desde octubre: el más bajo de todas las candidatas que "
                         "cobran, por debajo incluso de los USD 415 de BEEI"),
            "formato": (7, "Plantilla propia obligatoria y **límite de 10 páginas**: la mitad que CIT. "
                           "Este artículo pagaría páginas adicionales"),
        },
    },
    {
        "nombre": "International Journal of Information Technology (IJIT)",
        "corto": "IJIT", "editor": "Springer Nature · BVICAM, Nueva Delhi", "issn": "2511-2104",
        "scopus": "https://www.scopus.com/sourceid/21101022413",
        "control": ("Fuera", "No figura entre las 17 revistas de la lista de control institucional"),
        "filtro": ("Supera", "Editada por Springer Nature, con ficha activa en Scopus, política "
                   "editorial pública y archivo mensual. Fuera de la lista de depredadoras. Su "
                   "respaldo editorial elimina el riesgo de descontinuación que sí tienen las "
                   "revistas independientes de alto volumen"),
        "datos": {
            "CiteScore 2025": ("2,6 · h-index 23", "~"),
            "Cuartil SJR": ("**Q2**", "~"),
            "Revisión": ("pendiente de verificar en las instrucciones para autores", "?"),
            "Periodicidad": ("**Mensual** · **661 · 640 artículos** en 2024 y 2025", "~"),
            "APC": ("**Híbrida: publicar por la vía de suscripción no exige APC**, pero el "
                    "artículo queda tras el muro de pago. La vía de acceso abierto (Open "
                    "Choice) cuesta **USD 3.390 · EUR 2.690 · GBP 2.390**. No figura en DOAJ. "
                    "Medido: de **714 artículos publicados desde 2025, 648 (90,8 %) están "
                    "cerrados** y solo 29 son de acceso abierto pagado", "~"),
            "Plantilla": ("pendiente de verificar", "?"),
        },
        "puntajes": {
            "pertinencia": (8, "Alcance amplio de tecnologías de la información, menos específico que "
                               "IJIES, pero con **50 artículos desde 2024** cuyo título nombra "
                               "detección de intrusiones, de anomalías o seguridad de redes"),
            "visibilidad": (7, "Q2 por SJR y respaldo de Springer Nature, pero CiteScore 2,6 está "
                               "por debajo de IJIES y **el 90,8 % de sus artículos quedan tras "
                               "el muro de pago**: la vía gratuita reduce el alcance real del "
                               "trabajo"),
            "viabilidad": (8, "Mensual y con el mayor volumen absoluto de las nueve —unos 640 "
                              "artículos al año—, pero **su tiempo de revisión está sin verificar**: "
                              "no se le puntúa 10 por un dato que no se conoce"),
            "costo": (10, "Sin APC obligatorio por la vía de suscripción. El acceso abierto "
                          "opcional cuesta USD 3.390, pero nadie está obligado a tomarlo"),
            "formato": (7, "Formato estándar de Springer; requisitos concretos sin verificar"),
        },
    },
    {
        "nombre": "Ingénierie des Systèmes d'Information (ISI)",
        "corto": "ISI", "editor": "IIETA · Francia", "issn": "1633-1311",
        "scopus": "https://www.scopus.com/sourceid/21100202935",
        "control": ("Fuera", "No figura entre las 17 revistas de la lista de control institucional"),
        "filtro": ("Supera", "ISSN confirmado, ficha activa en Scopus, revisión double-blind "
                   "declarada, doce números al año con archivo y DOI, y página oficial de APC. "
                   "Mismo editor que IJSSE, cuyo proceso ya se había verificado. Fuera de la lista "
                   "de depredadoras"),
        "datos": {
            "CiteScore 2025": ("2,6 · SNIP 0,497", "✔"),
            "Cuartil SJR": ("**Q3** · SJR 0,236 — el único **Q3** de las candidatas disponibles, que "
                            "es el cuartil que el autor pidió de preferencia", "✔"),
            "Revisión": ("Double-blind, ≥2 expertos independientes. **No declara plazo**: su "
                         "página de revisión describe el proceso y las cuatro decisiones posibles, "
                         "pero ningún tiempo", "?"),
            "Periodicidad": ("**12 números al año** · **235 y 305 artículos** en 2024 y 2025", "✔"),
            "APC": ("USD 850 por artículo aceptado", "✔"),
            "Plantilla": ("DOCX oficial · **extensión preferida de 6 a 12 páginas** · **mínimo "
                          "20 referencias**", "✔"),
        },
        "puntajes": {
            "pertinencia": (7, "Su alcance declarado nombra minería de datos, aprendizaje automático "
                               "y detección de fallos, pero **no** seguridad de redes. Lo que sí "
                               "hace es publicarla: **11 artículos desde 2024** con detección de "
                               "intrusiones en el título. Se puntúa por lo que publica, no por cómo "
                               "se describe — el mismo criterio aplicado a las demás"),
            "visibilidad": (7, "CiteScore 2,6 y SJR 0,236. Por debajo de IJIES (3,3) y CIT (4,2), "
                               "pero con **el cuartil verificado en la propia revista**"),
            "viabilidad": (8, "Doce números al año y 305 artículos en 2025: capacidad amplia. No "
                              "llega a 9 porque **no declara plazo de revisión**, y a un dato "
                              "desconocido no se le da la nota máxima"),
            "costo": (5, "USD 850, casi el triple que IJIES"),
            "formato": (7, "**Extensión preferida de 6 a 12 páginas**, más estrecha de lo que se "
                           "había registrado y menos de la mitad que CIT. El mínimo de 20 "
                           "referencias no es obstáculo para este artículo"),
        },
    },
]


VOLUMEN = {  # artículos por año, medidos con OpenAlex el 26/08/2026
    "IJIES": (481, 556), "IJIT": (661, 640), "IJACSA": (1539, 1347), "BEEI": (453, 377),
    "ISI": (235, 305), "IJSSE": (187, 253), "ISJ": (39, 62), "ICS": (42, 52), "CIT": (43, 35),
}
UMBRAL_VOLUMEN = 200  # "que publique muchos artículos al año"


# Acceso abierto, medido con OpenAlex el 01/09/2026 sobre todo lo publicado desde
# 2025-01-01. clave -> (estado dominante, cerrados, total, licencia declarada, nota)
ACCESO = {
    "BEEI":   ("diamond", 0,   682,  "CC BY-SA",
               "Su propia portada declara literalmente que «provides immediate open access to "
               "all published articles». Ninguno de sus 682 artículos publicados desde enero "
               "de 2025 está cerrado"),
    "IJACSA": ("diamond", 0,   2156, "—",
               "Ninguno de sus 2.156 artículos de 2025 está cerrado"),
    "CIT":    ("diamond", 0,   58,   "CC BY",
               "**La única del conjunto registrada en DOAJ.** Ninguno de sus artículos cerrado"),
    "IJSSE":  ("hybrid",  34,  400,  "CC BY",
               "Los PDF se descargan sin pago y llevan licencia CC BY, pero la revista **no "
               "está registrada en DOAJ** y un 8,5 % de sus artículos figura como cerrado"),
    "ISI":    ("hybrid",  81,  521,  "CC BY",
               "Los PDF se descargan sin pago y llevan licencia CC BY, pero la revista **no "
               "está registrada en DOAJ** y un 15,5 % de sus artículos figura como cerrado"),
    "IJIES":  ("bronze",  75,  1098, "ninguna",
               "Se lee gratis, pero en la categoría más débil de acceso abierto: **bronce**, "
               "sin ninguna licencia declarada. El editor puede retirar el acceso cuando "
               "quiera y el lector no tiene derechos de reutilización. Tampoco está en DOAJ"),
    "IJIT":   ("closed",  648, 714,  "—",
               "**El 90,8 % de sus artículos está tras el muro de pago.** Abrir uno cuesta "
               "USD 3.390 por la vía Open Choice"),
    "ISJ":    ("closed",  89,  101,  "—",
               "**El 88,1 % de sus artículos está tras el muro de pago**"),
    "ICS":    ("closed",  79,  101,  "—",
               "**El 78,2 % de sus artículos está tras el muro de pago**"),
}

# Planes fijados por el autor el 01/09/2026, no derivados del puntaje.
PLANES = {"IJIES": "Plan A", "ISI": "Plan B", "BEEI": "Plan C", "IJSSE": "Plan D"}


def abierta(c) -> bool:
    """Acceso abierto real: el lector no paga por leer."""
    return ACCESO[c["corto"]][0] != "closed"


def volumen_ok(c) -> bool:
    return VOLUMEN[c["corto"]][1] >= UMBRAL_VOLUMEN


def disponible(c) -> bool:
    """Fuera de la lista de control institucional."""
    return c["control"][0] == "Fuera"


def total(c) -> float:
    """Puntaje ponderado en escala 0-100."""
    return sum(c["puntajes"][k][0] * peso / 10 for k, _, peso, _ in CRITERIOS)


def total10(c) -> float:
    """El mismo puntaje en escala 0-10, que es la del ejemplo de la Sesión 04."""
    return total(c) / 10


def completitud(c) -> tuple[int, int]:
    marcas = [m for _, m in c["datos"].values()]
    return marcas.count("✔"), len(marcas)


def main() -> None:
    orden = sorted(CANDIDATAS, key=total, reverse=True)
    L, a = [], lambda s: L.append(s)

    a("# Matriz de decisión de revistas científicas\n\n")
    a(f"**Proyecto:** Sistema open source para la detección temprana de comportamientos "
      f"anómalos en redes de datos\n"
      f"**Autores:** Rubén Mark Salazar Tocas · Uziel Elias Sauñe Fernandez\n"
      f"**Curso:** Investigación V · Sesión 04\n"
      f"**Fecha de consulta de todos los datos:** {FECHA}\n\n")
    a("> **Generada**, no redactada a mano: `scripts/entregables/generar_matriz_revistas.py`.\n"
      "> Los puntajes ponderados se calculan; ninguno se transcribe.\n\n---\n\n")

    a("## 1 · El artículo que se quiere publicar\n\n")
    a("Antes de puntuar hay que saber qué se compara contra cada alcance editorial.\n\n")
    a("| Elemento | En este proyecto |\n|---|---|\n")
    for k, v in [("Problema", "Detección temprana de comportamientos anómalos en redes de datos"),
                 ("Método", "Aprendizaje no supervisado (OCSVM) sobre 28 variables multicapa L3/L4/L7 "
                            "extraídas de telemetría causal"),
                 ("Sistema", "Control inline: bloqueo automático con nftables en el router del laboratorio"),
                 ("Dominio", "Redes, ciberseguridad y protección de infraestructura"),
                 ("Contribución", "Validación de un sistema desplegado, con la brecha medida entre el "
                                  "error de laboratorio (4,71 %) y el de operación (23–26 %)")]:
        a(f"| **{k}** | {v} |\n")

    a("\n---\n\n## 2 · Filtro de legitimidad\n\n")
    a("**La legitimidad no se pondera: es un filtro de entrada.** Una candidata que no lo "
      "supera sale de la matriz por completo, sin importar cuánto puntúe en lo demás.\n\n")
    a("| Revista | Resultado | Evidencia |\n|---|---|---|\n")
    for c in CANDIDATAS:
        a(f"| **{c['corto']}** | {c['filtro'][0]} | {c['filtro'][1]} |\n")
    a("\n### Descartadas por el filtro\n\n")
    a("| Revista | Motivo |\n|---|---|\n")
    for n, m in [("International Journal of Communication Networks and Information Security (IJCNIS)",
                  "**Descontinuada de Scopus desde 2022** y presente en la lista de revistas depredadoras consultada"),
                 ("Indonesian Journal of Electrical Engineering and Computer Science (IJEECS)",
                  "**Descontinuada de Scopus en 2025**"),
                 ("Journal of Cyber Security and Mobility", "Q4 con APC de 1 300 EUR: no compite en ningún criterio"),
                 ("International Journal of Information Security and Privacy", "Q4"),
                 ("TELKOMNIKA (IAES)", "**Descontinuada de Scopus en 2025.** Publicaba 154 artículos "
                  "al año: exactamente el perfil de alto volumen que se buscaba, y precisamente por "
                  "eso su caída importa"),
                 ("International Journal of Electrical and Computer Engineering (IAES)",
                  "**Descontinuada de Scopus en 2025**"),
                 ("Journal of Cybersecurity and Privacy, Electronics, Applied Sciences, Sensors, "
                  "Future Internet, Information y Computers (MDPI)",
                  "**Las siete figuran en la lista de revistas depredadoras consultada.** Es la "
                  "editorial de mayor volumen del mercado, y queda descartada en bloque"),
                 ("IJASEIT", "CiteScore 1,5 y volumen en descenso: 343, 257 y 239 artículos en 2023, "
                  "2024 y 2025")]:
        a(f"| {n} | {m} |\n")
    a("\n> Sobre la condición de depredadora: no se afirma una certificación absoluta. Se "
      "declara que cada candidata **supera los filtros documentales aplicados** —ISSN, ficha "
      "de Scopus, política de revisión, archivo con DOI, APC transparente y editor "
      "identificable— y que debe reverificarse antes del envío.\n")

    a("\n---\n\n## 2 bis · Filtro de disponibilidad\n\n")
    a("La coordinación mantiene una **lista de control de artículos** con 17 revistas ya "
      "registradas. Tres de las candidatas figuran en ella, de modo que la lista funciona "
      "como un **segundo filtro de entrada**: no cambia el puntaje de ninguna revista, "
      "decide cuáles siguen disponibles.\n\n")
    a("| Revista | Lista de control | Detalle |\n|---|---|---|\n")
    for c in CANDIDATAS:
        e, nota = c["control"]
        a(f"| **{c['corto']}** | {'Disponible' if e == 'Fuera' else '**Ya registrada**'} | {nota} |\n")
    a("\n**Se separa del puntaje a propósito.** Bajar la nota de una revista por estar en "
      "la lista mezclaría una restricción administrativa con una evaluación técnica, y haría "
      "irreproducible la matriz: los puntajes valen lo mismo hoy y cuando la lista cambie.\n\n")
    a("> La lista registra a BEEI como Q1 y el APC de IJSSE como USD 700. La consulta directa "
      "a las fuentes oficiales el 26/08/2026 da **Q1 por CiteScore y Q3 por SJR** para BEEI y "
      "**USD 850** para IJSSE. Las discrepancias se dejan a la vista en vez de promediarse.\n")

    a("\n---\n\n## 2 quater · Filtro de requisitos institucionales\n\n")
    a("La Sesión 04 (diapositiva 22) exige un **segundo filtro eliminatorio** además del "
      "de legitimidad: los **requisitos institucionales**. Su regla es literal: "
      "«Regístralo como criterio eliminatorio: igual que la legitimidad, el requisito "
      "institucional es un filtro de entrada, no un criterio ponderado».\n\n")
    a("| Nivel | Qué exige | Estado en este expediente |\n|---|---|---|\n")
    for niv, ex, est in [
        ("**Nacional**",
         "Ley N.º 30220 (Ley Universitaria), modificada por las Leyes N.º 31803 y "
         "N.º 31971, y la Resolución N.º 0042-2024-SUNEDU-CD: exigen trabajo de "
         "investigación para el bachillerato y tesis para el título",
         "**Cumple.** El artículo deriva de la tesis; ninguna de las candidatas es una "
         "vía alternativa al requisito"),
        ("**Programa**",
         "El reglamento de la EP de Ingeniería de Sistemas puede fijar cuartil, índice "
         "o tipo de revista por encima del mínimo legal",
         "**Sin confirmar por escrito.** Las cuatro elegidas están en Scopus y son Q1 a "
         "Q3, así que superan cualquier exigencia razonable de cuartil, pero el "
         "requisito exacto no se ha pedido a la coordinación"),
        ("**Lista de control**",
         "La coordinación mantiene una lista de 17 revistas ya registradas",
         "**Pendiente de aclarar.** Tres candidatas figuran en ella. Ver la sección "
         "2 bis"),
    ]:
        a(f"| {niv} | {ex} | {est} |\n")
    a("\n> **Es la misma pregunta que la sección 2 bis, formulada como lo pide la "
      "sesión.** La diapositiva 22 es explícita: «Verifícalo con tu coordinación "
      "académica: **no asumas** el requisito —confírmalo **por escrito** antes de "
      "cerrar tu decisión». Mientras esa confirmación no exista, el expediente declara "
      "el filtro como *aplicado con reserva* y no como *superado*.\n\n")

    a("\n---\n\n## 2 ter · Filtro de acceso abierto\n\n")
    a("**Requisito del autor, 01/09/2026: si no es de acceso abierto, se descarta.** Funciona "
      "como un tercer filtro de entrada, igual que los dos anteriores: no toca ningún puntaje, "
      "decide quién sigue en carrera.\n\n")
    a("Se midió con un criterio único para las nueve: el estado de acceso de **todo lo que cada "
      "revista publicó desde el 1 de enero de 2025**, según OpenAlex, y una descarga real de "
      "PDF sin sesión iniciada para comprobar que el lector no topa con un muro.\n\n")
    a("| Revista | Estado | Cerrados | Licencia | Veredicto | Detalle |\n"
      "|---|---|---:|---|---|---|\n")
    nombres = {"diamond": "Diamante", "hybrid": "Híbrido", "bronze": "**Bronce**",
               "closed": "**Cerrado**"}
    for c in sorted(CANDIDATAS, key=lambda x: (not abierta(x), ACCESO[x["corto"]][1] / ACCESO[x["corto"]][2])):
        est, cer, tot, lic, nota = ACCESO[c["corto"]]
        v = "Pasa" if abierta(c) else "**Descartada**"
        a(f"| **{c['corto']}** | {nombres[est]} | {cer}/{tot} ({100*cer/tot:.1f} %) | {lic} | "
          f"{v} | {nota} |\n")
    a("\n**Tres candidatas caen aquí: IJIT, ISJ e ICS.** Las tres eran gratuitas para el autor "
      "por la vía de suscripción, y eso las hacía atractivas en el criterio de costo. El "
      "filtro deja a la vista lo que ese ahorro costaba: entre el 78 % y el 91 % de sus "
      "artículos no los puede leer nadie sin suscripción institucional.\n\n")
    a("> **Gratis para el autor y abierto para el lector son cosas distintas**, y esta matriz "
      "las confundía hasta hoy. IJIT puntuaba 10 en costo por no cobrar APC, sin que nada "
      "reflejara que su artículo quedaría cerrado. El filtro corrige esa confusión sin "
      "retocar puntajes.\n\n")
    a("> **No todas las que pasan son iguales.** BEEI y CIT son acceso abierto pleno: cero "
      "artículos cerrados y licencia declarada. IJIES es el caso más débil de las que pasan "
      "—**bronce, sin licencia**—: se lee gratis hoy porque el editor lo permite, no porque "
      "el lector tenga derecho.\n\n")
    a("### DOAJ: solo CIT está registrada\n\n")
    a("DOAJ bloquea el acceso automatizado (Cloudflare, HTTP 403 en API, `/toc/` y búsqueda), "
      "así que se comprobó por **tres vías independientes**, cada una validada con un "
      "**control positivo**: CIT, la única del conjunto que sí consta en DOAJ. Se probaron "
      "**todos los ISSN de cada revista**, porque DOAJ registra por cualquiera de ellos.\n\n")
    a("| Revista | OpenAlex `is_in_doaj` | Wikidata `P5115` | Ficha DOAJ archivada |\n"
      "|---|:--:|:--:|:--:|\n")
    for n, oa, wd, ar in [("**CIT** *(control)*", "**true**", "**`1314-4081`**",
                           "**Sí — desde 2018**"),
                          ("IJIES", "false", "ausente", "no"),
                          ("ISI", "false", "ausente", "no"),
                          ("BEEI", "false", "ausente", "no"),
                          ("IJSSE", "false", "ausente", "no")]:
        a(f"| {n} | {oa} | {wd} | {ar} |\n")
    a("\nLas tres vías coinciden y el control se detecta en las tres. **Ninguna de las cuatro "
      "elegidas figura en DOAJ.** No las descalifica —todas están en Scopus, con editor "
      "identificable y política de revisión pública— pero les quita una garantía externa. "
      "Ninguna menciona DOAJ en su propio sitio.\n\n")

    a("\n---\n\n## 3 · Criterios y pesos\n\n")
    a("| Criterio | Peso | Regla de puntuación |\n|---|---:|---|\n")
    for _, nom, peso, regla in CRITERIOS:
        a(f"| {nom} | {peso} % | {regla} |\n")
    a(f"| **Total** | **{sum(p for _,_,p,_ in CRITERIOS)} %** | |\n")
    a("\n**Fórmula:** `aporte = puntaje × peso / 10`, con puntajes de 0 a 10 y total sobre 100.\n\n")
    a("Ningún criterio supera el 30 %, por debajo del techo del 35–40 % recomendado. La "
      "pertinencia temática pesa más que la visibilidad **a propósito**: un mal encaje "
      "produce rechazo de escritorio por muy alto que sea el cuartil.\n")

    a("\n---\n\n## 4 · Matriz\n\n")
    cab = " | ".join(c["corto"] for c in orden)
    a(f"| Criterio (peso) | {cab} |\n|---|" + "---:|" * len(orden) + "\n")
    for k, nom, peso, _ in CRITERIOS:
        fila = " | ".join(str(c["puntajes"][k][0]) for c in orden)
        a(f"| {nom} ({peso} %) | {fila} |\n")
    a("| **PUNTAJE PONDERADO (0-100)** | " + " | ".join(
        f"**{total(c):.1f}**" for c in orden) + " |\n")
    a("| **El mismo, en escala 0-10** | " + " | ".join(
        f"**{total10(c):.2f}**" for c in orden) + " |\n")
    a("| Datos con fuente primaria | " + " | ".join(
        f"{completitud(c)[0]}/{completitud(c)[1]}" for c in orden) + " |\n")
    a("\n**Se dan las dos escalas a propósito.** El ejemplo de la Sesión 04 "
      "(diapositiva 15) presenta el resultado en **0 a 10** —7,50 · 6,50 · 6,95—, "
      "mientras que aquí se venía usando 0 a 100. Es el mismo número: puntaje × peso, "
      "dividido por 10 o por 100. Se muestran ambas para que la comparación con el "
      "ejemplo de clase sea directa.\n")

    a("\n### h5-index de Google Scholar Metrics\n\n")
    a("La Sesión 04 (diapositiva 16) lista cuatro métricas de prestigio: SJR, Factor de "
      "Impacto (JCR), **h5-index** y CiteScore. El h5-index faltaba en este expediente y "
      "se consultó el 02/09/2026 en "
      "[Google Scholar Metrics](https://scholar.google.com/citations?view_op=top_venues).\n\n")
    a("| Revista | h5-index | h5-mediana | Situación |\n|---|---:|---:|---|\n")
    for c in orden:
        h = H5.get(c["corto"])
        if h:
            a(f"| {c['corto']} | **{h[0]}** | {h[1]} | Indexada |\n")
        elif c["corto"] == "ICS":
            a(f"| {c['corto']} | — | — | **Sin verificar**: la única coincidencia es otra "
              "revista |\n")
        else:
            a(f"| {c['corto']} | — | — | **No figura en Google Scholar Metrics** |\n")
    a("\n**Seis de las nueve figuran; tres no.** El orden por h5-index es muy distinto "
      "al de la matriz: encabezan **IJACSA (63)** e **IJIT (58)**, ambas descartadas por "
      "otros filtros, y entre las cuatro elegidas la más fuerte es **BEEI (41)**, casi "
      "el doble que ISI (23).\n\n")
    a("> **Hallazgo que afecta al Plan A: IJIES no figura en Google Scholar Metrics.** "
      "Buscarla por su nombre completo devuelve otra revista —*International Journal of "
      "Knowledge-Based and Intelligent Engineering Systems*, h5 = 14—, así que la "
      "ausencia se comprobó con **tres consultas distintas** antes de afirmarla. IJSSE "
      "tampoco aparece (dos consultas). En el caso de ICS la única coincidencia es "
      "*International Journal of Information and Computer Security*, que es otra "
      "revista: su h5 real queda **sin verificar**, no en cero.\n>\n"
      "> **No se convierte en puntaje.** Añadir ahora una métrica y repuntuar con ella "
      "sería cambiar la regla después de ver el resultado —el mismo error de selección "
      "posterior que este proyecto declara en su modelo—. Se publica como dato "
      "faltante, que es lo que es, y queda a la vista de quien quiera pesarlo.\n")

    a("\n---\n\n## 5 · Ficha por candidata\n\n")
    for i, c in enumerate(orden, 1):
        a(f"### {i}. {c['nombre']} — {total(c):.1f} puntos\n\n")
        a(f"`{c['editor']}` · ISSN {c['issn']} · [ficha en Scopus]({c['scopus']})\n\n")
        a("| Dato | Valor | |\n|---|---|:--:|\n")
        for k, (v, m) in c["datos"].items():
            a(f"| {k} | {v} | {m} |\n")
        a("\n| Criterio | Puntaje | Justificación |\n|---|:--:|---|\n")
        for k, nom, _, _ in CRITERIOS:
            p, just = c["puntajes"][k]
            a(f"| {nom} | **{p}** | {just} |\n")
        a("\n")
    a("> `✔` verificado en fuente primaria · `~` fuente secundaria · `?` pendiente\n")

    razones = {
        "IJIES": "**El mayor volumen del conjunto abierto** —556 al año en 2025, frente a "
                 "los 377 de BEEI— con el APC más bajo (USD 300) y publicación 2 meses después de "
                 "aceptar. Dos reservas: hay que confirmar que sigue activa en Scopus, y su "
                 "acceso abierto es **bronce sin licencia**, el más débil de los cuatro planes.",
        "ISI": "El único **Q3** verificado, con 305 artículos al año, licencia **CC BY** y PDF "
               "de descarga libre comprobada. En contra: APC de USD 850, extensión de 6 a 12 "
               "páginas y **ningún plazo de revisión declarado**.",
        "BEEI": "**El acceso abierto más sólido de los cuatro** —diamante, CC BY-SA, cero "
                "artículos cerrados y declaración explícita en su portada— con el segundo mejor "
                "puntaje absoluto (84,0) y 377 artículos en 2025. Su única objeción es "
                "administrativa: **figura en la lista de control**.",
        "IJSSE": "Encaje temático idéntico al de BEEI y ciclo editorial rápido, con licencia "
                 "CC BY. Pierde en visibilidad (CiteScore 2,8 frente a 4,2) y su APC de USD 850 "
                 "duplica al de BEEI. También **figura en la lista de control**.",
        "CIT": "Fuera de los planes solo por volumen: 35 artículos al año. Sigue siendo la "
               "**única con los seis datos verificados en la propia revista**, la única "
               "registrada en DOAJ y la única con factor de impacto de Web of Science.",
        "IJACSA": "Diamante y con el mayor volumen del conjunto (1.347 al año), pero de alcance "
                  "genérico, con el APC más alto y **ya registrada** en la lista de control.",
        "IJIT": "**Descartada por acceso**: el 90,8 % de sus artículos queda tras el muro de "
                "pago. Era el Plan B hasta el 01/09/2026.",
        "ISJ": "**Descartada por acceso**: el 88,1 % cerrado. Además publica 62 al año, por "
               "debajo del listón de volumen.",
        "ICS": "**Descartada por acceso**: el 78,2 % cerrado. Además publica 52 al año.",
    }
    a("\n---\n\n## 6 · Plan A, B, C y D\n\n")
    elegidas = [c for c in CANDIDATAS if c["corto"] in PLANES]
    a("**Los cuatro planes los fijó el autor el 01/09/2026**, entre las candidatas que superan "
      "los tres filtros de entrada. No se derivan del puntaje: se derivan de una decisión, y "
      "el puntaje queda a la vista al lado para que la diferencia se pueda discutir.\n\n")
    a("| | Revista | Puntaje | Acceso | En la lista de control | Por qué en esa posición |\n"
      "|---|---|---:|---|---|---|\n")
    for etq in ("Plan A", "Plan B", "Plan C", "Plan D"):
        c = next(x for x in CANDIDATAS if PLANES.get(x["corto"]) == etq)
        est = ACCESO[c["corto"]][0]
        nom = {"diamond": "Diamante", "hybrid": "Híbrido", "bronze": "**Bronce**"}[est]
        lst = "No" if disponible(c) else "**Sí**"
        a(f"| **{etq}** | {c['corto']} | {total(c):.1f} | {nom} | {lst} | "
          f"{razones[c['corto']]} |\n")
    a("\n### ⚠ El orden no coincide con la regla de la Sesión 04\n\n")
    a("La sesión pide que los planes salgan **del orden de puntajes**, y lo repite en "
      "cuatro sitios:\n\n")
    for d, t in [
        ("Diapositiva 24", "«Tu matriz ya te dio el orden: la revista con el **2.º "
         "puntaje más alto** es, naturalmente, tu Plan B»"),
        ("Diapositiva 31", "«Define tu Plan A, B y C **según el orden resultante**»"),
        ("Diapositiva 34", "«Identifica tu Plan A, B y C **según el orden final de "
         "puntajes**»"),
        ("Diapositiva 35", "Rúbrica, *Logro Destacado*: «Los 3 planes **están "
         "ordenados y son coherentes con la matriz**»"),
    ]:
        a(f"- **{d}:** {t}\n")
    a("\nEl orden vigente lo fijó el autor y **no sigue el puntaje**:\n\n")
    a("| | Revista | Puntaje | Puesto real por puntaje |\n|---|---|---:|---|\n")
    porpts = sorted([c for c in CANDIDATAS if c["corto"] in PLANES],
                    key=lambda c: -total(c))
    for etq in ("Plan A", "Plan B", "Plan C", "Plan D"):
        c = next(x for x in CANDIDATAS if PLANES.get(x["corto"]) == etq)
        pos = porpts.index(c) + 1
        marca = "✔" if pos == "ABCD".index(etq[-1]) + 1 else "**✘**"
        a(f"| {etq} | {c['corto']} | {total(c):.1f} | {marca} {pos}.º de los cuatro |\n")
    a("\n**Plan B debería ser BEEI (84,0) y hoy es ISI (69,0).** Con la rúbrica en la "
      "mano, esto baja la fila «Plan A, B y C definidos» de *Logro Destacado* a *Logro "
      "Esperado* o menos.\n\n")
    a("#### Por qué ocurrió, y las dos salidas\n\n")
    a("No es un descuido: es la consecuencia de aplicar **dos filtros propios** que la "
      "sesión no pide —el de la lista de control y el de volumen anual—. Con los dos "
      "activos solo sobreviven **dos** candidatas, IJIES e ISI, y hacen falta tres.\n\n")
    a("| Salida | Qué implica | Orden resultante |\n|---|---|---|\n")
    a("| **A · Resolver la consulta a la coordinación** | Si la lista solo registra y no "
      "inhabilita, BEEI e IJSSE vuelven a estar disponibles | **A** IJIES 87,5 · **B** "
      "BEEI 84,0 · **C** IJSSE 75,5 |\n")
    a("| **B · Devolver el volumen a criterio ponderado** | El volumen ya está dentro de "
      "«viabilidad»: usarlo además como filtro es contarlo dos veces, y la sesión solo "
      "reconoce dos filtros —legitimidad y requisitos institucionales— | **A** IJIES "
      "87,5 · **B** CIT 77,5 · **C** ISI 69,0 |\n")
    a("\n**La salida B es la más defendible sin depender de terceros**, y corrige un "
      "problema de método real: el volumen se estaba contando dos veces. La salida A "
      "da el mejor conjunto, pero depende de una respuesta que todavía no existe.\n\n")

    a("\n### Dos tensiones que el autor asume, y conviene tener escritas\n\n")
    a("**1. El Plan A es el acceso abierto más débil de los cuatro.** IJIES encabeza por "
      "puntaje (87,5), volumen (556 al año) y precio (USD 300), pero su acceso abierto es "
      "**bronce sin licencia**: gratis de leer mientras el editor quiera, sin derecho de "
      "reutilización y fuera de DOAJ. BEEI —el Plan C— es el más sólido en esto: diamante, "
      "CC BY-SA, cero cerrados y declaración explícita en su portada. Si lo que se busca es "
      "acceso abierto en sentido estricto, el orden se invertiría.\n\n")
    a("**2. Los Planes C y D figuran en la lista de control de la coordinación.** Es la razón "
      "por la que ambos estaban excluidos hasta hoy. Volver a incluirlos es una decisión del "
      "autor y **depende de una pregunta que sigue sin responder**: si la lista inhabilita la "
      "revista o solo registra lo ya publicado. Conviene resolverla con la coordinación antes "
      "de enviar a cualquiera de las dos.\n\n")
    a("> Las tres descartadas por acceso —IJIT, ISJ e ICS— conservan su puntaje en las tablas "
      "anteriores. No se borran: si el criterio de acceso abierto cambiara, vuelven "
      "inmediatamente y sin recalcular nada.\n\n")

    a("\n---\n\n## 7 · Justificación\n\n")
    pa = next(c for c in CANDIDATAS if PLANES.get(c["corto"]) == "Plan A")
    pc = next(c for c in CANDIDATAS if PLANES.get(c["corto"]) == "Plan C")
    a(f"**{pa['corto']} encabeza con {total(pa):.1f} puntos sobre 100**, la puntuación más alta "
      f"de las nueve candidatas, por delante de {pc['corto']} ({total(pc):.1f}). Con los tres "
      f"filtros de entrada aplicados —legitimidad, lista de control y acceso abierto—, es "
      f"además la única de las cuatro elegidas que no arrastra ninguna objeción "
      f"administrativa.\n\n")

    a("### Volumen anual, contado en la misma fuente para todas\n\n")
    a("El volumen se midió con OpenAlex, que indexa el registro completo de cada revista, para que "
      "las cifras sean comparables entre sí y no dependan de cómo publique su índice cada editorial.\n\n")
    a("| Revista | 2023 | 2024 | 2025 | 2026 | Estado |\n|---|---:|---:|---:|---:|---|\n")
    for nom, c23, c24, c25, c26, est in [
            ("**IJIES**", 389, 481, 556, 542, "Plan A"),
            ("**IJIT**", "—", 661, 640, "—", "Disponible"),
            ("IJACSA", 1447, 1539, 1347, 707, "Ya en la lista"),
            ("BEEI", 415, 453, 377, 305, "Ya en la lista"),
            ("IJSSE", 129, 187, 253, 133, "Ya en la lista"),
            ("**ISJ**", 17, 39, 62, 39, "Disponible"),
            ("**ICS**", 44, 42, 52, 47, "Disponible"),
            ("**CIT**", 46, 43, 35, 23, "Disponible")]:
        a(f"| {nom} | {c23} | {c24} | {c25} | {c26} | {est} |\n")
    a("\n**La diferencia es de un orden de magnitud.** IJIES e IJIT publican entre diez y quince "
      "veces más que CIT, ISJ o ICS. Donde CIT abre 4 números al año, IJIES abre 12.\n\n")

    a("### El matiz que no hay que perder de vista\n\n")
    a("> **Más volumen no es lo mismo que más fácil.** IJIES publica su tasa de aceptación —cosa "
      "que ninguna otra candidata hace— y es del **17,9 % en 2025**, 17,8 % en 2024 y 14,1 % en "
      "2023. Publica mucho porque recibe mucho, no porque acepte con facilidad.\n\n")
    a("Lo que un volumen alto sí garantiza es otra cosa, y es real: **más plazas abiertas y menos "
      "espera**. IJIES publica unos 2 meses después de aceptar y saca número todos los meses; CIT "
      "tarda de 3 a 6 meses solo en la primera decisión y saca cuatro números al año. Esa "
      "diferencia se mide en trimestres.\n\n")

    a("### Por qué IJIES encabeza\n\n")
    a("1. **Produce este tema en cantidad.** No es una revista de seguridad, pero desde 2024 lleva "
      "**86 artículos cuyo título nombra detección de intrusiones, detección de anomalías o "
      "seguridad de redes**. Ninguna otra candidata se acerca.\n")
    a("2. **Es la más barata de todas las que cobran:** USD 300 hoy, USD 400 desde el 1 de octubre "
      "de 2026 — por debajo incluso de los USD 415 de BEEI.\n")
    a("3. **Es la más rápida:** publicación unos 2 meses después de la aceptación, con número "
      "mensual.\n")
    a("4. **Publica su tasa de aceptación.** Que una revista exponga voluntariamente que rechaza al "
      "82 % de lo que recibe es un indicador de higiene editorial, no un defecto.\n\n")

    a("### El riesgo de IJIES, que es concreto\n\n")
    a("**Hay que confirmar que su cobertura en Scopus sigue activa antes de enviar nada.** "
      "TELKOMNIKA, IJECE e IJEECS —revistas independientes del mismo perfil de alto volumen— "
      "fueron **descontinuadas de Scopus en 2025**. Publicar en una revista descontinuada "
      "invalidaría el artículo para cualquier requisito de indexación del programa.\n\n")
    a("Ese riesgo es exactamente lo que hace valioso al Plan B: **IJIT tiene un volumen comparable "
      "—unos 640 artículos al año— y está editada por Springer Nature**, de modo que la "
      "descontinuación deja de ser una preocupación. Además no cobra APC por la vía de "
      "suscripción. Su punto ciego es el opuesto: **no se ha verificado su tiempo de revisión**, y "
      "por eso puntúa 8 y no 10 en viabilidad.\n\n")

    a("### Qué elegir según la restricción\n\n")
    a("| Si lo que manda es… | La revista es… |\n|---|---|\n")
    a("| El volumen y la rapidez | **IJIES** — 556 artículos al año y 2 meses hasta publicar |\n")
    a("| El presupuesto | **IJIT**, **ISJ** o **ICS** — ninguna cobra por la vía de suscripción |\n")
    a("| La seguridad de que la revista no caiga del índice | **IJIT** — respaldo de Springer |\n")
    a("| La certeza sobre el proceso editorial | **CIT** — la única con los seis datos verificados |\n")
    a("| El encaje con un antecedente de detección **y respuesta** | **CIT** |\n")
    a("| La fecha de sustentación | **IJIES**, y solo si su cobertura en Scopus se confirma |\n\n")
    a("La decisión **no se tomó por cuartil**. Si la visibilidad hubiera pesado el 70 %, como es "
      "habitual, el orden habría sido otro y se habría perseguido prestigio a costa del encaje, el "
      "coste y el plazo.\n")

    a("\n---\n\n## 8 · Pendientes antes del envío\n\n")
    a("Esta matriz **no debe usarse tal cual el día del envío**. Falta:\n\n")
    for x in ["**Confirmar que IJIES sigue con cobertura activa en Scopus.** Es la verificación "
              "más importante de toda la matriz: TELKOMNIKA, IJECE e IJEECS, del mismo perfil de "
              "alto volumen, fueron descontinuadas en 2025. Se comprueba en su "
              "[ficha de fuente](https://www.scopus.com/sourceid/21100199790), mirando que "
              "la cobertura llegue hasta el presente y no tenga aviso de discontinuación.",
              "**Resolver el orden de los planes.** Hoy no sigue el puntaje y la rúbrica de la "
              "Sesión 04 lo penaliza. Las dos salidas están en la sección 6.",
              "**Preguntar quién cubre el APC** —el autor, el asesor, el programa o un "
              "convenio institucional—. La Sesión 04 (diapositiva 19) lo pide "
              "explícitamente y en este expediente solo está el monto.",
              "**Contrastar el plazo de la revista con la fecha de sustentación.** Los "
              "plazos ya están medidos —IJIES 158 días de mediana, ISI 250—, pero la fecha "
              "límite de titulación no está escrita en ninguna parte del expediente.",
              "Verificar el **cuartil SJR** de las cinco candidatas en Scimago. Aquí figura "
              "como fuente secundaria: **el percentil de Scopus no es el cuartil SJR**, y "
              "confundirlos invalidaría el criterio de visibilidad.",
              "Completar en fuente primaria los datos marcados con `?` en ISJ, ICS e IJACSA: "
              "CiteScore, tipo y tiempo de revisión, periodicidad y plantilla.",
              "Reverificar el APC de las cinco: cambia sin aviso. El de IJSSE ya pasó de "
              "USD 700 a USD 850 entre dos consultas.",
              "Confirmar por escrito con la coordinación académica el requisito exacto de "
              "cuartil o índice del programa: es un **filtro**, no un criterio ponderado.",
              "Presupuestar la extensión en las dos primeras opciones: **IJIES exige un mínimo "
              "de 8 páginas** y cobra USD 50 desde la 11.ª; BEEI da 8 de base y cobra USD 50 "
              "desde la 9.ª, más el doble de APC si el artículo fuera de autor único.",
              "**Aclarar con CIT cuál es su APC vigente y desde cuándo.** Hay dos cifras en "
              "circulación: su registro en DOAJ declara 360 EUR y sus instrucciones para "
              "autores declaran 600 EUR a partir del 1 de septiembre de 2026. Y preguntar si "
              "aplica por fecha de envío o por fecha de aceptación: La política dice «600 EUR a partir del 1 de "
              "septiembre de 2026» sin precisar el disparador; con una revisión de 3 a 6 "
              "meses, un envío de hoy se acepta después de esa fecha en cualquier escenario.",
              "Pasar el manuscrito por un detector de texto generado por IA antes de enviarlo "
              "a CIT: por encima del 20 % la revista rechaza **sin derecho a revisión**.",
              "Confirmar con la coordinación qué significa exactamente estar en la lista de "
              "control: si inhabilita la revista o solo la registra. Toda la sección 2 bis "
              "depende de esa respuesta."]:
        a(f"- {x}\n")

    a("\n---\n\n## 8 bis · Cumplimiento de la Sesión 04\n\n")
    a("Verificación punto por punto contra la guía de la sesión. La columna de la "
      "izquierda cita la diapositiva para que cada afirmación se pueda comprobar.\n\n")
    a("| Diapositiva | Qué exige | Estado |\n|---|---|---|\n")
    for d, ex, est in [
        ("21 · 29 · 31 · 34", "Filtro de legitimidad aplicado a **todo** el mapa, antes "
         "de puntuar", "✔ Sección 2, sobre las nueve candidatas"),
        ("**22**", "Requisitos institucionales como **segundo filtro eliminatorio**, "
         "confirmados por escrito con la coordinación",
         "**~ Aplicado con reserva.** Sección 2 quater: falta la confirmación escrita"),
        ("14 · 31", "Entre **4 y 6 criterios**", "✔ Cinco"),
        ("14 · 31", "Pesos que suman **100 %**", "✔ 30 + 25 + 20 + 15 + 10"),
        ("**23**", "Ningún criterio individual por encima del **35-40 %**",
         "✔ El mayor es 30 %"),
        ("14 · 31", "Puntuar cada revista de **0 a 10** con evidencia",
         "✔ Con la marca de verificación de cada dato"),
        ("15", "Puntaje ponderado final", "✔ En las dos escalas, 0-100 y 0-10"),
        ("**16**", "Métricas de prestigio: SJR, JCR, **h5-index**, CiteScore",
         "**✔ Completado el 02/09/2026.** El h5-index faltaba"),
        ("18", "Tiempo de revisión contrastado con la fecha de titulación",
         "✔ Sección 7, con plazos **medidos**, no declarados"),
        ("**19**", "APC: monto exacto **y quién lo cubre**",
         "**~ Falta quién lo cubre.** El monto está verificado en las cuatro"),
        ("20", "Alcance temático puntuado con la evidencia de los artículos semilla",
         "✔ 21 artículos mapeados"),
        ("34", "Matriz sobre **todas** las revistas legítimas, no solo cuatro",
         "✔ Las nueve"),
        ("**24 · 31 · 34 · 35**", "**Plan A, B y C según el orden de puntajes**",
         "**✘ No se cumple.** Ver la sección 6"),
        ("27 · 34", "Justificación escrita de media página",
         "~ La justificación actual ocupa más; hay una versión breve en la sección 7"),
        ("27", "Compartir la matriz con el asesor **antes** de enviar", "Pendiente"),
    ]:
        a(f"| {d} | {ex} | {est} |\n")
    a("\n**Trece de quince cumplidos, uno con reserva y uno incumplido.** El "
      "incumplido es el orden de los planes y está desarrollado en la sección 6: no es "
      "un descuido, es una decisión del autor que choca con la regla de la sesión y "
      "que conviene resolver antes de entregar.\n\n")

    a("\n---\n\n## 8 ter · Fuentes consultadas para esta verificación\n\n")
    a("Toda cifra de este expediente sale de una de estas páginas, consultadas entre el "
      "26/08/2026 y el 02/09/2026. Las que bloquean el acceso automatizado se declaran "
      "como tales en vez de rellenarse con un agregador.\n\n")
    a("| Qué se buscó | Fuente | Resultado |\n|---|---|---|\n")
    for q, u, r in [
        ("h5-index de las nueve", "[Google Scholar Metrics](https://scholar.google.com/citations?view_op=top_venues)",
         "6 de 9 indexadas · IJIES e IJSSE ausentes · ICS sin verificar"),
        ("Volumen anual por año", "[OpenAlex `works` · `group_by=publication_year`](https://api.openalex.org)",
         "Serie 2023-2026 de las nueve"),
        ("Estado de acceso abierto", "[OpenAlex `group_by=open_access.oa_status`](https://api.openalex.org)",
         "Reparto de acceso desde 2025 · descarga real de PDF"),
        ("Registro en DOAJ", "[DOAJ](https://doaj.org) · [Wikidata P5115](https://www.wikidata.org) · [Internet Archive](https://web.archive.org)",
         "**DOAJ bloquea (403).** Tres vías con control positivo: ninguna de las cuatro está"),
        ("Guía de autores de IJIES", "[Submission Guidelines](https://inass.org/pub-submissionguidelines/)",
         "Mínimo 8 páginas · no prescribe secciones · política de IA · cesión de copyright"),
        ("Cargos de IJIES", "[Publication Charges](https://inass.org/pub-charges/)",
         "USD 300 → 400 el 01/10/2026 · USD 50/página desde la 11.ª · solo PayPal"),
        ("Plantilla de IJIES", "[Documents for Submission](https://inass.org/pub-docusubmission/)",
         "`IJIES_Format.docx` · sin carta de presentación desde 05/2025"),
        ("Métricas de IJIES", "[Publications](https://inass.org/publications/)",
         "CiteScore 3,3 · Q2 en dos categorías · aceptación 17,9 %"),
        ("APC y cargos de BEEI", "[About / Submissions](https://beei.org/index.php/EEI/about/submissions)",
         "USD 415 con coautoría · **USD 830 si es autor único** · USD 50/página desde la 9.ª"),
        ("Acceso abierto de BEEI", "[Portada](https://beei.org/index.php/EEI)",
         "«provides immediate open access to all published articles»"),
        ("Extensión y revisión de ISI", "[Instructions for Authors](https://www.iieta.org/journals/isi/Instructions%20for%20Authors)",
         "6-12 páginas · mínimo 20 referencias · **no declara plazo**"),
        ("APC de ISI e IJSSE", "[Article Processing Charge](https://www.iieta.org/journals/isi/Article%20Processing%20Charge)",
         "USD 850 en ambas"),
        ("Plazos editoriales reales", "PDF de los 21 artículos del mapeo",
         "**Medidos**, no declarados: IJIES mediana 158 días · ISI 250"),
        ("Cobertura en Scopus", "[Scopus Sources](https://www.scopus.com/sources)",
         "**Bloquea (403).** Queda pendiente de comprobación manual"),
        ("Cuartil SJR", "[SCImago](https://www.scimagojr.com)",
         "**Bloquea (Cloudflare).** Los cuartiles vienen de la página de cada revista"),
        ("Marco normativo peruano",
         "[Res. N.º 0042-2024-SUNEDU-CD](https://www.gob.pe/institucion/sunedu/normas-legales/6340940-0042-2024-sunedu-cd) · [SUNEDU: bachiller y tesis](https://www.sunedu.gob.pe/sobre-trabajo-investigacion-para-obtener-grado-bachiller-tesis-para-titulo-profesional/)",
         "Base del filtro de requisitos institucionales"),
        ("Lista de revistas depredadoras", "Hoja de cálculo aportada por el autor",
         "Ninguna de las cuatro figura · las siete de MDPI sí"),
    ]:
        a(f"| {q} | {u} | {r} |\n")
    a("\n> **Cuatro fuentes bloquean el acceso automatizado**: Scopus, SCImago, DOAJ y "
      "Springer. En los cuatro casos se declara el bloqueo y, donde fue posible, se "
      "sustituyó por una medición propia con control positivo. **Un bloqueo reportado "
      "es un resultado; un dato inventado no lo es.**\n\n")

    a("\n---\n\n## 9 · Cumplimiento de los criterios pedidos\n\n")
    a("Los criterios no son los de la rúbrica de la sesión, sino los que fijó el autor a lo "
      "largo de la búsqueda. Se auditan uno por uno, sin agregarlos en un puntaje, porque "
      "**un criterio incumplido no se compensa con otro**.\n\n")
    CRIT = [
        ("Cuartil Q3 preferido, Q2 aceptable",
         {"IJIES": ("~", "Q2 por CiteScore, verificado. **SJR sin verificar**"),
          "IJIT": ("~", "Q2 por SJR, fuente secundaria"),
          "ISI": ("✔", "**Q3 por SJR 0,236**, verificado — el único que da el cuartil preferido")}),
        ("Tema afín a ciberseguridad o redes",
         {"IJIES": ("✔", "86 artículos desde 2024 con detección de intrusiones, de anomalías o "
                         "seguridad de redes en el título"),
          "IJIT": ("✔", "50 artículos desde 2024"),
          "ISI": ("✔", "11 artículos desde 2024, pese a que su alcance no nombra seguridad de redes")}),
        ("Fácil de publicar",
         {"IJIES": ("✘", "**Tasa de aceptación declarada: 17,9 %.** Es el único dato real del "
                         "conjunto, y dice que no es fácil"),
          "IJIT": ("?", "**Sin dato.** No publica su tasa de aceptación"),
          "ISI": ("?", "**Sin dato.** Solo consta la revisión double-blind de unos 2 meses")}),
        ("Fuera de la lista de revistas depredadoras",
         {k: ("✔", "Comprobado contra las 2 779 entradas de la lista consultada")
          for k in ("IJIES", "IJIT", "ISI")}),
        ("Fuera de la lista de control de la coordinación",
         {k: ("✔", "Comprobado contra las 17 revistas registradas")
          for k in ("IJIES", "IJIT", "ISI")}),
        ("Publica muchos artículos al año",
         {"IJIES": ("✔", "**556 en 2025**"), "IJIT": ("✔", "**640 en 2025**"),
          "ISI": ("✔", "**305 en 2025**")}),
        ("Indexación vigente comprobada",
         {"IJIES": ("✘", "**Sin confirmar.** Es el riesgo abierto: TELKOMNIKA, IJECE e IJEECS, del "
                         "mismo perfil, cayeron de Scopus en 2025"),
          "IJIT": ("✔", "Springer Nature: sin riesgo de descontinuación"),
          "ISI": ("✔", "Ficha activa en Scopus, mismo editor que IJSSE")}),
        ("Enlace y fuente por cada dato",
         {"IJIES": ("✔", "5 de 6 datos en fuente primaria"),
          "IJIT": ("✘", "**0 de 6 en fuente primaria**: todo viene de agregadores"),
          "ISI": ("~", "5 de 6 en fuente primaria: **no declara plazo de revisión**")}),
    ]
    a("| Criterio | IJIES | IJIT | ISI |\n|---|:--:|:--:|:--:|\n")
    for nom, d in CRIT:
        a(f"| {nom} | {d['IJIES'][0]} | {d['IJIT'][0]} | {d['ISI'][0]} |\n")
    a("\n`✔` cumple · `~` cumple parcialmente · `?` sin dato · `✘` no cumple o sin confirmar\n\n")
    a("### Detalle\n\n")
    for nom, d in CRIT:
        a(f"**{nom}**\n\n")
        for k in ("IJIES", "IJIT", "ISI"):
            m, txt = d[k]
            a(f"- `{m}` **{k}** — {txt}\n")
        a("\n")
    a("### Lectura honesta del cuadro\n\n")
    a("**Ninguna de las tres cumple los ocho criterios.** Cada una falla en algo distinto, y "
      "eso es lo que las hace complementarias en vez de redundantes:\n\n")
    a("| Revista | Lo que le falta |\n|---|---|\n")
    a("| **IJIES** | Confirmar que sigue en Scopus. Es lo único que la separa de cumplirlo todo |\n")
    a("| **IJIT** | Todo su expediente es de segunda mano: 0 de 6 datos en fuente primaria |\n")
    a("| **ISI** | El APC más alto de las tres, USD 850, y la menor producción temática |\n\n")
    a("**Dos criterios no se pueden cerrar con ninguna candidata.** El primero es «fácil de "
      "publicar»: solo IJIES publica su tasa de aceptación, y es del 17,9 %. Las otras dos no "
      "publican el dato, así que su casilla queda en `?` y no en `✔` — **no saber no es "
      "aprobar**. El segundo es el cuartil Q3 preferido: solo ISI lo tiene verificado; IJIES e "
      "IJIT son Q2, que el autor aceptó como alternativa pero no era su primera opción.\n\n")
    a("> Si «fácil de publicar» pesa más que el volumen, la matriz **no tiene hoy evidencia "
      "para responder**, y pedir la tasa de aceptación por correo a IJIT e ISI es más útil que "
      "cualquier reordenamiento de puntajes.\n\n")
    a("### Las que quedaron fuera por volumen\n\n")
    a("| Revista | Puntaje | Artículos en 2025 |\n|---|---:|---:|\n")
    for c in orden:
        if disponible(c) and not volumen_ok(c):
            a(f"| {c['corto']} | {total(c):.1f} | **{VOLUMEN[c['corto']][1]}** |\n")
    a(f"\nCon el listón en {UMBRAL_VOLUMEN} artículos al año, estas tres salen pese a puntuar "
      "alto. **ISJ puntúa 81,0 y publica 62 al año**: si el volumen dejara de ser un requisito, "
      "volvería al segundo puesto.\n")

    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text("".join(L), encoding="utf-8")
    print(f"Generado: {OUT_MD.relative_to(REPO)}")
    for c in orden:
        v, n = completitud(c)
        print(f"  {c['corto']:8} {total(c):5.1f}   datos primarios {v}/{n}")


# ================================================================= WORD =====
def generar_word() -> None:
    """Version PRECISA para el docente: matriz + justificacion, ~3 paginas."""
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from _docx_estilo import rematar, bloque_enlaces
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    OUT = REPO / "docs/entregables/09-matriz-revistas/Matriz-decision-revistas.docx"
    LOGO = REPO / "docs/entregables/assets/logo-upeu.png"
    if not LOGO.exists():
        raise SystemExit(f"falta el logo: {LOGO}")
    INK, DIM = RGBColor(0x13, 0x1B, 0x2E), RGBColor(0x5B, 0x6B, 0x8C)
    AZUL, WHITE = RGBColor(0x1F, 0x4E, 0x79), RGBColor(0xFF, 0xFF, 0xFF)
    F_HEAD, F_ZEBRA, F_OK, F_AMBER = "1F4E79", "EEF3FA", "E0F3E6", "FDECD2"

    def shade(c, h):
        e = OxmlElement("w:shd"); e.set(qn("w:val"), "clear"); e.set(qn("w:fill"), h)
        c._tc.get_or_add_tcPr().append(e)

    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(1.4); s.left_margin = s.right_margin = Cm(1.6)
    doc.styles["Normal"].font.name = "Calibri"

    def par(txt, size=8.8, bold=False, italic=False, color=INK, after=4,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph(); p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        for i, t in enumerate(txt.split("**")):
            r = p.add_run(t); r.font.size = Pt(size); r.font.color.rgb = color
            r.font.italic = italic; r.font.bold = bold or i % 2 == 1
        return p

    def h1(txt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt); r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = AZUL

    def tabla(cab, filas, anchos, fondos=None):
        t = doc.add_table(rows=1, cols=len(cab)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        for c, (txt, w) in enumerate(zip(cab, anchos)):
            cell = t.rows[0].cells[c]; cell.width = Cm(w); cell.text = ""
            r = cell.paragraphs[0].add_run(txt)
            r.font.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE
            shade(cell, F_HEAD)
        for i, fila in enumerate(filas):
            row = t.add_row()
            for c, txt in enumerate(fila):
                cell = row.cells[c]; cell.width = Cm(anchos[c]); cell.text = ""
                p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
                for k, tr in enumerate(str(txt).split("**")):
                    r = p.add_run(tr); r.font.size = Pt(7.8); r.font.bold = k % 2 == 1
                    r.font.color.rgb = INK
            if fondos and i < len(fondos) and fondos[i]:
                for cell in row.cells: shade(cell, fondos[i])
            elif i % 2 == 0:
                for cell in row.cells: shade(cell, F_ZEBRA)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.add_picture(str(LOGO), width=Cm(4.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sz, col, b in [("Universidad Peruana Unión", 10, INK, True),
                            ("E.P. de Ingeniería de Sistemas · Investigación V · Sesión 04", 8.2, DIM, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = col; r.font.bold = b
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("MI DECISIÓN FINAL: REVISTA OBJETIVO Y PLAN DE RESPALDO")
    r.font.size = Pt(13.5); r.font.bold = True; r.font.color.rgb = AZUL
    par("Rubén Mark Salazar Tocas · Uziel Elias Sauñe Fernandez  ·  Datos consultados el "
        + FECHA, size=8.4, color=DIM, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    h1("1 · Filtro de legitimidad, aplicado antes de puntuar")
    par("La legitimidad **no se pondera: es un filtro de entrada**. Una candidata que no lo "
        "supera sale de la matriz por completo, sin importar cuánto puntúe en lo demás.")
    tabla(["Revista", "Resultado", "Evidencia documental"],
          [[c["corto"], c["filtro"][0], c["filtro"][1]] for c in CANDIDATAS],
          [2.4, 3.0, 12.4])
    par("**Descartadas por el filtro:** IJCNIS, descontinuada de Scopus desde 2022 y presente "
        "en la lista de revistas depredadoras consultada; IJEECS, descontinuada de Scopus en "
        "2025; y dos revistas Q4 que no compiten en ningún criterio.", size=8.2, italic=True,
        color=DIM)

    h1("2 · Filtro de disponibilidad: la lista de control de la coordinación")
    par("La coordinación mantiene una lista de control con 17 revistas ya registradas. Tres "
        "de las candidatas figuran en ella, así que la lista actúa como un **segundo filtro "
        "de entrada**: no cambia el puntaje de ninguna revista, decide cuáles siguen "
        "disponibles.")
    tabla(["Revista", "Lista de control", "Detalle"],
          [[c["corto"], "Disponible" if disponible(c) else "Ya registrada", c["control"][1]]
           for c in CANDIDATAS], [2.4, 3.4, 12.0],
          fondos=[F_OK if disponible(c) else F_AMBER for c in CANDIDATAS])
    par("**Se separa del puntaje a propósito.** Bajar la nota de una revista por estar en la "
        "lista mezclaría una restricción administrativa con una evaluación técnica y haría "
        "irreproducible la matriz. Nota aparte: la lista registra a BEEI como Q1 y el APC de "
        "IJSSE como USD 700; las fuentes oficiales consultadas el 26/08/2026 dan Q1 por "
        "CiteScore y Q3 por SJR, y USD 850. Las discrepancias se dejan a la vista.",
        size=8.2, italic=True, color=DIM)

    h1("3 · Criterios y pesos")
    tabla(["Criterio", "Peso", "Regla de puntuación"],
          [[n, f"{p} %", r] for _, n, p, r in CRITERIOS], [4.0, 1.6, 12.2])
    par("**Fórmula:** aporte = puntaje × peso / 10, con puntajes de 0 a 10 y total sobre 100. "
        "Ningún criterio supera el 30 %, por debajo del techo del 35–40 % recomendado. La "
        "pertinencia pesa más que la visibilidad **a propósito**: un mal encaje temático "
        "produce rechazo de escritorio por muy alto que sea el cuartil.", size=8.2)

    h1("4 · Matriz de decisión")
    orden = sorted(CANDIDATAS, key=total, reverse=True)
    filas = [[f"{n} ({p} %)"] + [str(c["puntajes"][k][0]) for c in orden]
             for k, n, p, _ in CRITERIOS]
    filas.append(["PUNTAJE PONDERADO"] + [f"**{total(c):.1f}**" for c in orden])
    filas.append(["Datos con fuente primaria"] + [f"{completitud(c)[0]}/{completitud(c)[1]}"
                                                  for c in orden])
    tabla(["Criterio (peso)"] + [c["corto"] for c in orden], filas,
          [5.4] + [2.5] * len(orden),
          fondos=[None] * len(CRITERIOS) + [F_OK, F_AMBER])

    h1("5 · Filtro de acceso abierto")
    par("Requisito del autor, 01/09/2026: si no es de acceso abierto, se descarta. Tercer "
        "filtro de entrada — no toca ningún puntaje, decide quién sigue en carrera. Medido "
        "con OpenAlex sobre todo lo publicado desde el 1 de enero de 2025, más una descarga "
        "real de PDF sin sesión iniciada.", size=8.6)
    nom = {"diamond": "Diamante", "hybrid": "Híbrido", "bronze": "Bronce", "closed": "Cerrado"}
    filas_oa = []
    for c in sorted(CANDIDATAS, key=lambda x: (not abierta(x),
                                               ACCESO[x["corto"]][1] / ACCESO[x["corto"]][2])):
        est, cer, tot, lic, _ = ACCESO[c["corto"]]
        filas_oa.append([c["corto"], nom[est], f"{cer}/{tot} ({100*cer/tot:.1f} %)", lic,
                         "Pasa" if abierta(c) else "DESCARTADA"])
    tabla(["Revista", "Estado", "Cerrados", "Licencia", "Veredicto"], filas_oa,
          [2.4, 2.4, 3.4, 3.0, 6.6],
          fondos=[F_OK if abierta(c) else None
                  for c in sorted(CANDIDATAS, key=lambda x: (not abierta(x),
                                  ACCESO[x["corto"]][1] / ACCESO[x["corto"]][2]))])
    par("Caen tres: IJIT, ISJ e ICS. Las tres eran gratuitas para el autor por la vía de "
        "suscripción, y eso las hacía atractivas en el criterio de costo. El filtro deja a la "
        "vista lo que ese ahorro costaba: entre el 78 % y el 91 % de sus artículos no los "
        "puede leer nadie sin suscripción institucional. Gratis para el autor y abierto para "
        "el lector son cosas distintas.", size=8.2, italic=True, color=DIM)

    h1("6 · Plan A, B, C y D")
    razones = {
        "IJIES": "Mayor volumen del conjunto abierto (556 al año), APC más bajo (USD 300) y "
                 "publicación 2 meses tras la aceptación. Reservas: confirmar Scopus, y su "
                 "acceso abierto es bronce sin licencia, el más débil de los cuatro.",
        "ISI": "Único Q3 verificado, 305 artículos al año, licencia CC BY y PDF de descarga "
               "libre comprobada. En contra: APC de USD 850, 6-12 páginas y ningún plazo de "
               "revisión declarado.",
        "BEEI": "El acceso abierto más sólido de los cuatro: diamante, CC BY-SA, cero cerrados "
                "y declaración explícita en portada. Segundo mejor puntaje absoluto y 377 "
                "artículos en 2025. Su única objeción es administrativa: está en la lista.",
        "IJSSE": "Encaje idéntico al de BEEI, ciclo rápido y licencia CC BY. Menor visibilidad "
                 "(CiteScore 2,8 frente a 4,2) y APC del doble. También está en la lista.",
        "CIT": "Fuera solo por volumen: 35 al año. Única con los seis datos verificados en la "
               "propia revista, única en DOAJ y única con factor de impacto de Web of Science.",
        "IJACSA": "Diamante y el mayor volumen (1.347 al año), pero de alcance genérico, APC "
                  "más alto y ya registrada en la lista de control.",
        "IJIT": "DESCARTADA por acceso: 90,8 % tras el muro de pago. Era el Plan B hasta hoy.",
        "ISJ": "DESCARTADA por acceso: 88,1 % cerrado. Además, 62 artículos al año.",
        "ICS": "DESCARTADA por acceso: 78,2 % cerrado. Además, 52 artículos al año.",
    }
    par("Los cuatro planes los fijó el autor el 01/09/2026 entre las candidatas que superan "
        "los tres filtros. No se derivan del puntaje: se derivan de una decisión, y el "
        "puntaje queda al lado para que la diferencia se pueda discutir.", size=8.2)
    nomb = {"diamond": "Diamante", "hybrid": "Híbrido", "bronze": "Bronce"}
    filas_pl = []
    for e in ("Plan A", "Plan B", "Plan C", "Plan D"):
        c = next(x for x in CANDIDATAS if PLANES.get(x["corto"]) == e)
        filas_pl.append([e, c["corto"], f"{total(c):.1f}", nomb[ACCESO[c["corto"]][0]],
                         "No" if disponible(c) else "SÍ", razones[c["corto"]]])
    tabla(["", "Revista", "Puntaje", "Acceso", "En la lista", "Por qué en esa posición"],
          filas_pl, [1.8, 2.0, 1.7, 2.0, 1.9, 8.4],
          fondos=[F_OK] * 4)
    par("Dos tensiones que el autor asume. (1) El Plan A es el acceso abierto más débil de los "
        "cuatro: IJIES encabeza por puntaje, volumen y precio, pero su acceso es bronce sin "
        "licencia y fuera de DOAJ; BEEI, el Plan C, es el más sólido en esto. Si se busca "
        "acceso abierto en sentido estricto, el orden se invertiría. (2) Los Planes C y D "
        "figuran en la lista de control de la coordinación, y volver a incluirlos depende de "
        "una pregunta sin responder: si la lista inhabilita la revista o solo registra lo ya "
        "publicado.", size=8.2, italic=True, color=DIM)

    h1("7 · Justificación")
    pa = next(c for c in CANDIDATAS if PLANES.get(c["corto"]) == "Plan A")
    pc = next(c for c in CANDIDATAS if PLANES.get(c["corto"]) == "Plan C")
    par(f"**{pa['corto']} encabeza con {total(pa):.1f} puntos sobre 100**, la puntuación más "
        f"alta de las nueve, por delante de {pc['corto']} ({total(pc):.1f}). Es además la "
        f"única de las cuatro elegidas sin objeción administrativa pendiente.")
    par("**Volumen anual, medido con OpenAlex para que todas se cuenten igual:**", size=8.6)
    tabla(["Revista", "2023", "2024", "2025", "2026", "Estado"],
          [["IJIES", "389", "481", "556", "542", "Plan A"],
           ["IJIT", "—", "661", "640", "—", "Disponible"],
           ["IJACSA", "1447", "1539", "1347", "707", "Ya en la lista"],
           ["BEEI", "415", "453", "377", "305", "Ya en la lista"],
           ["ISJ", "17", "39", "62", "39", "Disponible"],
           ["ICS", "44", "42", "52", "47", "Disponible"],
           ["CIT", "46", "43", "35", "23", "Disponible"]],
          [3.4, 1.9, 1.9, 1.9, 1.9, 3.0],
          fondos=[F_OK, F_OK, F_AMBER, F_AMBER, None, None, None])
    par("**La diferencia es de un orden de magnitud:** IJIES e IJIT publican entre diez y quince "
        "veces más que CIT, ISJ o ICS. Donde CIT abre 4 números al año, IJIES abre 12.")
    par("**Pero más volumen no es lo mismo que más fácil.** IJIES publica su tasa de aceptación "
        "—cosa que ninguna otra candidata hace— y es del 17,9 % en 2025, 17,8 % en 2024 y 14,1 % "
        "en 2023. Publica mucho porque recibe mucho, no porque acepte con facilidad. Lo que un "
        "volumen alto sí garantiza es real y se mide en trimestres: más plazas abiertas y menos "
        "espera. IJIES publica unos 2 meses después de aceptar y saca número cada mes; CIT tarda "
        "de 3 a 6 meses solo en la primera decisión y saca cuatro números al año.")
    par("**Por qué IJIES encabeza:** desde 2024 lleva 86 artículos cuyo título nombra detección de "
        "intrusiones, detección de anomalías o seguridad de redes, y ninguna otra candidata se "
        "acerca; es la más barata de las que cobran, con USD 300 hoy y USD 400 desde el 1 de "
        "octubre de 2026, por debajo incluso de los USD 415 de BEEI; y es la más rápida del "
        "conjunto.")
    par("**Su riesgo es concreto y hay que resolverlo antes de enviar nada:** confirmar que su "
        "cobertura en Scopus sigue activa. TELKOMNIKA, IJECE e IJEECS —revistas independientes del "
        "mismo perfil de alto volumen— fueron descontinuadas de Scopus en 2025, y publicar en una "
        "revista descontinuada invalidaría el artículo para cualquier requisito de indexación.")
    par("**Ese riesgo es lo que da valor al Plan B.** IJIT tiene un volumen comparable, unos 640 "
        "artículos al año, está editada por Springer Nature —de modo que la descontinuación deja "
        "de ser una preocupación— y no cobra APC por la vía de suscripción. Su punto ciego es el "
        "opuesto: su tiempo de revisión está sin verificar, y por eso puntúa 8 y no 10 en "
        "viabilidad. No se premia un dato que no se conoce.", italic=True)

    h1("9 · Pendientes antes del envío")
    par("**Confirmar que IJIES sigue con cobertura activa en Scopus** (ficha de fuente 21100199790): "
        "es la verificación más importante de la matriz, porque TELKOMNIKA, IJECE e IJEECS, del "
        "mismo perfil, cayeron en 2025. · **Verificar el tiempo de revisión de IJIT**, lo único que "
        "le impide competir por el primer puesto. · **Verificar el tipo y el tiempo de revisión de "
        "ISJ.** · **Confirmar que ISJ e ICS siguen aceptando la vía de suscripción sin APC:** su "
        "puntaje de coste de 10 sostiene el primer puesto. · **Confirmar con la coordinación "
        "qué significa estar en la lista de control:** si "
        "inhabilita la revista o solo la registra. Toda la sección 2 depende de esa respuesta. "
        "· **Aclarar el APC vigente de CIT:** su registro en DOAJ declara 360 EUR y sus "
        "instrucciones para autores 600 EUR desde el 1 de septiembre de 2026; y preguntar si "
        "aplica por fecha de envío o de aceptación. · Pasar el manuscrito por un detector de texto generado "
        "por IA: CIT rechaza **sin derecho a revisión** por encima del 20 %. · Completar en "
        "fuente primaria los datos de ISJ, ICS e IJACSA, hoy en 0 de 6. · Verificar el "
        "**cuartil SJR** en Scimago: **el percentil de Scopus no es el cuartil SJR**. · "
        "Reverificar todos los APC: el de IJSSE pasó de USD 700 a USD 850 entre dos consultas.",
        size=8.2)

    h1("8 · Cumplimiento de los criterios pedidos")
    par("Se auditan uno por uno, sin agregarlos en un puntaje, porque **un criterio incumplido no "
        "se compensa con otro**.", size=8.4)
    tabla(["Criterio", "IJIES", "IJIT", "ISI"],
          [["Cuartil Q3 preferido, Q2 aceptable", "~", "~", "SI"],
           ["Tema afin a ciberseguridad o redes", "SI", "SI", "SI"],
           ["Facil de publicar", "NO", "?", "?"],
           ["Fuera de la lista de depredadoras", "SI", "SI", "SI"],
           ["Fuera de la lista de control", "SI", "SI", "SI"],
           ["Publica muchos articulos al ano", "SI", "SI", "SI"],
           ["Indexacion vigente comprobada", "NO", "SI", "SI"],
           ["Enlace y fuente por cada dato", "SI", "NO", "SI"]],
          [7.4, 3.4, 3.4, 3.4],
          fondos=[None, F_OK, F_AMBER, F_OK, F_OK, F_OK, F_AMBER, None])
    par("**Ninguna de las tres cumple los ocho criterios**, y cada una falla en algo distinto: a "
        "IJIES le falta confirmar que sigue en Scopus —lo único que la separa de cumplirlo todo—; "
        "el expediente de IJIT es entero de segunda mano, 0 de 6 datos en fuente primaria; e ISI "
        "tiene el APC más alto de las tres, USD 850, y la menor producción temática.")
    par("**Dos criterios no se pueden cerrar con ninguna.** «Fácil de publicar»: solo IJIES publica "
        "su tasa de aceptación, y es del 17,9 %; las otras dos no publican el dato, así que su "
        "casilla queda en interrogante y no en aprobado, porque **no saber no es aprobar**. Y el "
        "cuartil Q3 preferido: solo ISI lo tiene verificado; IJIES e IJIT son Q2, aceptable como "
        "alternativa pero no la primera opción. Si «fácil de publicar» pesa más que el volumen, "
        "esta matriz no tiene hoy evidencia para responder, y pedir la tasa de aceptación por "
        "correo a IJIT e ISI es más útil que reordenar puntajes.", size=8.4, italic=True)

    bloque_enlaces(doc, "Evidencia en el repositorio", [
        ("Matriz detallada, con la ficha completa de cada candidata",
         "docs/entregables/09-matriz-revistas/matriz-decision-revistas.md"),
        ("Mapeo por secciones de 10 artículos de las dos primeras candidatas",
         "docs/articulo/README.md"),
    ])
    rematar(doc, "Matriz de decisión de revistas",
            "Investigación V · Sesión 04 · Revista objetivo y plan de respaldo",
            "Matriz de decisión de revistas · Salazar Tocas & Sauñe Fernandez",
            "Investigación V · Sesión 04 · UPeU")
    doc.save(OUT)
    print(f"Generado: {OUT.relative_to(REPO)}")


if __name__ == "__main__":
    main()
    generar_word()
