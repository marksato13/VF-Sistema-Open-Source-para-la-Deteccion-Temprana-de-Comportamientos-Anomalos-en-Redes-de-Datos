#!/usr/bin/env python3
"""Word del informe de evaluacion critica (Sesion 01, momento CREA).

El .md de la carpeta es la fuente DETALLADA. Este Word es la version PRECISA
para el docente, con la estructura exacta que pide la consigna y dentro del
limite de 2-4 paginas.

    python3 scripts/entregables/generar_evaluacion_critica_word.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from docx.shared import Cm, Pt, RGBColor

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from _docx_estilo import rematar, bloque_enlaces

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "docs/entregables/assets/logo-upeu.png"
GRAF = REPO / "docs/entregables/graficas"
OUT = REPO / "docs/entregables/01-evaluacion-critica/Informe-evaluacion-critica.docx"

INK, DIM = RGBColor(0x13, 0x1B, 0x2E), RGBColor(0x5B, 0x6B, 0x8C)
ACCENT, WHITE = RGBColor(0x1F, 0x4E, 0x79), RGBColor(0xFF, 0xFF, 0xFF)
F_HEAD, F_ZEBRA, F_OK, F_AMBER, F_RED = "1F4E79", "EEF3FA", "E0F3E6", "FDECD2", "FBE3E1"


_MARCAS = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _tramos(txt):
    """Parte el texto en (contenido, negrita, cursiva, monoespaciado).

    Los helpers solo entendian **negrita**; *cursiva* y `codigo` se escribian
    tal cual y el lector veia los asteriscos y las comillas.
    """
    for t in _MARCAS.split(str(txt)):
        if not t:
            continue
        limpio = lambda x: x.replace("**", "").replace("*", "").replace("`", "")
        if t.startswith("**") and t.endswith("**") and len(t) > 4:
            yield limpio(t[2:-2]), True, False, False
        elif t.startswith("*") and t.endswith("*") and len(t) > 2:
            yield limpio(t[1:-1]), False, True, False
        elif t.startswith("`") and t.endswith("`") and len(t) > 2:
            yield limpio(t[1:-1]), False, False, True
        else:
            yield t, False, False, False


def shade(cell, hx):
    el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hx)
    cell._tc.get_or_add_tcPr().append(el)


def h1(doc, n, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{n}  {_MARCAS.sub(lambda m: m.group(0).strip("*`"), txt)}")
    r.font.size = Pt(12.5); r.font.bold = True; r.font.color.rgb = ACCENT
    return p


def par(doc, txt, size=9.2, italic=False, color=INK, after=5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after); p.alignment = align
    for t, _b, _i, _m in _tramos(txt):
        r = p.add_run(t)
        r.font.size = Pt(size); r.font.color.rgb = color
        r.font.italic = italic; r.font.bold = _b
        r.font.italic = r.font.italic or _i
        r.font.name = "Consolas" if _m else r.font.name
    return p


def tabla(doc, cab, filas, anchos, fondos=None):
    t = doc.add_table(rows=1, cols=len(cab))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for c, (txt, w) in enumerate(zip(cab, anchos)):
        cell = t.rows[0].cells[c]; cell.width = Cm(w); cell.text = ""
        r = cell.paragraphs[0].add_run(txt)
        r.font.bold = True; r.font.size = Pt(8.4); r.font.color.rgb = WHITE
        shade(cell, F_HEAD)
    for i, fila in enumerate(filas):
        row = t.add_row()
        for c, txt in enumerate(fila):
            cell = row.cells[c]; cell.width = Cm(anchos[c]); cell.text = ""
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            for tr, _b, _i, _m in _tramos(txt):
                r = p.add_run(tr)
                r.font.size = Pt(8.2); r.font.bold = _b
                r.font.italic = r.font.italic or _i
                r.font.name = "Consolas" if _m else r.font.name; r.font.color.rgb = INK
            if fondos and fondos[i]:
                shade(cell, fondos[i])
            elif i % 2 == 0:
                shade(cell, F_ZEBRA)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def figura(doc, nombre, pie, ancho=13.0):
    f = GRAF / nombre
    if not f.exists():
        raise SystemExit(f"falta la figura: {f}")
    doc.add_picture(str(f), width=Cm(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
    par(doc, pie, size=7.8, italic=True, color=DIM, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)


def main() -> None:
    if not LOGO.exists():
        raise SystemExit(f"falta el logo: {LOGO}")
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(1.4); s.left_margin = s.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"

    doc.add_picture(str(LOGO), width=Cm(4.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sz, col, b in [("Universidad Peruana Unión", 10, INK, True),
                            ("Facultad de Ingeniería y Arquitectura · E.P. de Ingeniería de Sistemas", 8.4, DIM, False),
                            ("Investigación V · Sesión 01", 8.4, DIM, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.color.rgb = col; r.font.bold = b

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("INFORME DE EVALUACIÓN CRÍTICA DE RESULTADOS")
    r.font.size = Pt(14.5); r.font.bold = True; r.font.color.rgb = ACCENT
    par(doc, "Sistema open source para la detección temprana de comportamientos anómalos "
             "en redes de datos\nRubén Mark Salazar Tocas · Uziel Elias Sauñe Fernandez\nAsesores: Ing. Nemias Saboya Ríos · Ing. Fernando Manuel Asin Gómez",
        size=9, color=DIM, after=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    par(doc, "**Entregable de la Sesión 01.** Este documento es el informe de resultados y evaluación crítica; "
             "el plan prospectivo de la Sesión 02 está en `07-plan-de-validacion/plan-de-validacion-de-resultados.md`.",
        size=8.2, color=DIM, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------------------------------------------------------------- 1
    h1(doc, "1 ·", "Lo que está listo: qué se abordó de manera objetiva")
    par(doc, "El producto es un sistema desplegado que detecta comportamiento anómalo y "
             "**bloquea la IP ofensora en el propio router**, validado sobre tráfico real de "
             "laboratorio. Toda cifra procede de artefactos verificables; ninguna se transcribe "
             "a mano. La evaluación se ordena por los **criterios de validez y confiabilidad** "
             "trabajados en la sesión.")
    tabla(doc, ["Criterio de la sesión", "Cómo se abordó en este proyecto", "Estado"], [
        ["**Validez interna**\n*¿El resultado se debe a lo que decimos, o algo lo contamina?*",
         "**Nuestros datos** se reparten por corrida completa, no al azar: así el modelo no "
         "puede reconocer la corrida en vez del ataque. Detectamos una fuga real y la "
         "corregimos.", "Abordada"],
        ["**Validez externa**\n*¿Funcionaría fuera de nuestro laboratorio?*",
         "**Nuestro modelo** se probó con validación cruzada de 5 pliegues y 58 corridas en la "
         "red real. Falta una jornada nueva: hasta tenerla **no afirmamos** que funcione en "
         "otra red.", "Parcial"],
        ["**Confiabilidad**\n*Si repetimos la medición, ¿da lo mismo?*",
         "**Nuestro sistema es determinista**: 10 ajustes repetidos dieron el mismo modelo y el "
         "mismo umbral. Dos pases completos de validación dieron resultados equivalentes.",
         "Abordada"],
        ["**Validez de constructo**\n*Lo que medimos, ¿mide lo que decimos?*",
         "**Nuestras 28 variables** están definidas una por una con su fórmula, ventana y "
         "fuente, generadas desde el código que corre. Ninguna se describe de memoria.",
         "Abordada"],
        ["**α de Cronbach / Kappa**\n*¿Por qué no las usamos?*",
         "**Nuestro producto no es un cuestionario** ni tiene jueces clasificando, así que "
         "ninguna de las dos aplica. El Alfa se calculará sobre el SUS, que sí tiene 10 ítems.",
         "Justificado"],
    ], [4.0, 9.2, 4.2])
    par(doc, "**ISO/IEC 25010 — características evaluadas con evidencia:** *adecuación funcional* "
             "(detecta y bloquea, 88,8 % de detección), *eficiencia de desempeño* (bloqueo en "
             "mediana de 8,0 s), *confiabilidad* (cero caídas en 58 corridas, determinismo "
             "verificado) y *seguridad* (trazabilidad por SHA-256). **Sin evidencia todavía**: "
             "*usabilidad* —el SUS está pendiente—, *compatibilidad*, *mantenibilidad* y "
             "*portabilidad*.", size=8.8)
    par(doc, "**Ya resuelto en agosto**, sin capturar datos nuevos ni reentrenar: intervalos de "
             "Wilson en toda proporción y McNemar con corrección de Holm; ablación por capas y "
             "comparación de 14 contra 28 variables; diccionario de las 28 variables; dataset, "
             "manifiesto y 7 modelos publicados con sus checksums; datasheet, model card y system "
             "card; declaración de la selección posterior; validación cruzada y bootstrap.",
        size=8.8)
    par(doc, "**Resultados medidos:**", size=9.2)
    tabla(doc, ["Resultado medido", "Valor obtenido", "Cómo se verificó"], [
        ["**Capacidad discriminante**", "**ROC-AUC 0,974**", "Re-puntuando el modelo congelado"],
        ["**Detección de ataques genuinos**", "**88,8 %** [83,0 – 92,8]", "Evaluación bloqueada de un solo paso, con intervalo de Wilson"],
        ["**Respuesta en tiempo real**", "Bloqueo en mediana de **8,0 s**; rango 6,1–13,7 s (n = 8)", "58 corridas con el motor activo; 8 bloqueos con tiempo observable"],
        ["**Disponibilidad**", "**Cero caídas registradas**; 55 de 58 corridas con verificación explícita", "Registro de servicios antes y después de cada corrida"],
        ["**Aporte de las variables multicapa**", "De **66,5 % a 88,8 %** de detección, p < 0,001", "Ablación por capas con McNemar exacto"],
        ["**Superioridad del modelo elegido**", "Las **6 comparaciones** del OCSVM son significativas", "McNemar + corrección de Holm sobre 21 pares"],
        ["**Reproducibilidad**", "Dataset, manifiesto y **los 7 modelos** publicados y verificables", "sha256sum -c · licencias MIT y CC BY 4.0"],
    ], [4.4, 5.4, 7.6])

    # ---------------------------------------------------------------- 2
    h1(doc, "2 ·", "Lo que NO está listo")
    par(doc, "Solo lo que sigue abierto hoy. **Lo ya resuelto está en la sección 1 y no se repite "
             "aquí.** Las cinco están medidas, no supuestas.")
    tabla(doc, ["Qué falta", "Evidencia de que falta", "Gravedad"], [
        ["**El sistema bloquea tráfico legítimo pesado**", "En laboratorio se equivoca el 4,71 % de las veces; en operación, 25,81 % y 22,97 %. Una transferencia legítima de 200 Mbit/s bloqueó a un cliente real durante 120 s", "Crítica"],
        ["**El modelo se eligió mirando el examen final**", "Se compararon 7 candidatos sobre el mismo conjunto de prueba y se escogió el mejor. El 88,8 % es un máximo, no una estimación limpia", "Crítica"],
        ["**Nadie ha usado el panel salvo el equipo**", "El instrumento SUS está preparado, pero el archivo de respuestas tiene 0 filas: no se ha aplicado", "Alta"],
        ["**No se sabe si funciona en otra jornada**", "Los 44 perfiles aparecen en las tres particiones. Falta una captura nueva que el modelo no haya visto", "Alta"],
        ["**Faltan 4 escenarios legítimos del jurado**", "SSH, SCP/SFTP, backup y actualizaciones no están en el dataset", "Media"],
    ], [5.2, 9.0, 2.2], fondos=[F_RED, F_RED, F_AMBER, F_AMBER, F_ZEBRA])

    # ---------------------------------------------------------------- 3
    h1(doc, "3 ·", "Cómo se va a abordar")
    par(doc, "Una acción por cada fila de la sección 2, en el mismo orden. Todas caben en el "
             "laboratorio actual: **ninguna exige equipo ni presupuesto nuevo**.")
    tabla(doc, ["Qué falta", "Qué se hará", "¿Lo resuelve del todo?"], [
        ["Bloquea tráfico legítimo pesado", "Reentrenar incluyendo la transferencia de 200 Mbit/s como tráfico **normal** y repetir las 29 corridas de validación", "Sí, si el error baja. Si no baja, queda declarado como límite del sistema"],
        ["Modelo elegido mirando el examen", "Recolectar una jornada nueva que el modelo no vea nunca, y medir sobre ella", "Sí. Es la única corrección real; no se puede arreglar escribiendo"],
        ["Nadie ha usado el panel", "Sesión de 2 h con 5–8 evaluadores usando el instrumento SUS ya preparado", "Sí"],
        ["No sabemos si funciona otro día", "La misma jornada nueva de la fila 2 sirve para las dos cosas", "Sí"],
        ["Faltan 4 escenarios legítimos", "Una campaña F1 más: SSH, SCP/SFTP, backup y actualizaciones", "Sí"],
    ], [4.4, 7.4, 4.8])
    # -------------------------------------------------------- cronograma
    h1(doc, "4 ·", "En qué tiempo: cronograma comprometido")
    par(doc, "*Propuesta del equipo del 2 de septiembre de 2026, pendiente del visto bueno de los "
             "asesores.* **La fecha que ordena todo es el 30 de septiembre de 2026:** IJIES sube su "
             "APC de USD 300 a USD 400 el 1 de octubre y el cargo lo asume la Universidad Peruana "
             "Unión, así que enviar antes ahorra USD 100 institucionales. Con la mediana medida de "
             "la revista —41 días a la primera decisión, 158 hasta publicar—, un envío el 28 de "
             "septiembre proyecta decisión hacia el 8 de noviembre de 2026 y publicación hacia "
             "marzo de 2027.")
    tabla(doc, ["Fecha", "Pendiente", "Cómo se aborda", "Responsable", "Estado"], [
        ["**vie 4 sep 2026**", "Declarar la selección posterior en la tesis", "Párrafo en metodología, enlazado a la *model card*", "Salazar", "PLANIFICADA"],
        ["**mié 9 sep 2026**", "**Validación con usuarios (SUS)**", "Sesión de 2 h con 5–8 evaluadores; instrumento ya preparado", "Salazar · Sauñe", "PLANIFICADA"],
        ["**sáb 19 sep 2026**", "Escenarios legítimos faltantes", "Campaña F1: SSH, SCP/SFTP, backup y actualizaciones", "Sauñe", "PLANIFICADA"],
        ["**mié 23 sep 2026**", "Juicio experto (3 evaluadores)", "Rúbrica de pertinencia, ya con los resultados del SUS a la vista", "Salazar · asesores", "PLANIFICADA"],
        ["**lun 28 sep 2026**", "**Envío del artículo a IJIES**", "Manuscrito en `IJIES_Format.docx`, 8 a 10 páginas", "Salazar · Sauñe", "PLANIFICADA"],
        ["**sáb 10 oct 2026**", "**Recalibrar con tráfico pesado y repetir F6**", "Reentrenar con `iperf-tcp 200M` como normalidad; repetir las 29 corridas", "Salazar", "PLANIFICADA"],
        ["**sáb 24 oct 2026**", "*Holdout* temporal externo", "Campaña completa en fecha distinta, sin reutilizar episodios", "Salazar · Sauñe", "PLANIFICADA"],
    ], [2.9, 4.0, 5.6, 2.4, 2.3])
    par(doc, "**El SUS va primero porque es el único cero absoluto que queda**: cero en la ficha de "
             "auditoría, cero en el eje de pertinencia y `D-18` en el registro. Cuesta dos horas y "
             "sube la ficha de 82,4 % a 88,2 %. El juicio experto va después a propósito: los "
             "expertos juzgan mejor con los resultados de los usuarios delante.", italic=True)
    par(doc, "**El artículo no espera a los tres últimos.** La sección de resultados se escribe con "
             "lo que ya está bloqueado —modelo congelado, evaluación de un solo paso y 58 corridas "
             "de F6—: ese resultado está completo y no va a cambiar. Si la recalibración se retrasa, "
             "el artículo sale igual con la limitación declarada, que es como debe salir de todos "
             "modos. **La fecha de sustentación no la fija el equipo**, pero ningún pendiente de "
             "esta lista la condiciona más allá del 24 de octubre de 2026.", italic=True)

    # ------------------------------------------------------- limitaciones
    h1(doc, "5 ·", "Amenazas a la validez (*Threats to Validity*)")
    par(doc, "La sesión advierte que esta sección es **obligatoria en las revistas indexadas de "
             "Ingeniería de Software**. Se declara aquí y se trasladará al artículo.", size=8.8)
    tabla(doc, ["Tipo de amenaza", "Amenaza concreta en este trabajo", "Cómo se mitiga o declara"], [
        ["**Validez interna**", "El modelo se eligió observando el conjunto de prueba: **selección "
         "posterior**, la misma familia de error que el HARKing", "Declarada en la *model card* "
         "antes de cualquier métrica. La corrección exige evaluación nueva y reservada"],
        ["**Validez externa**", "Entorno artificial de laboratorio, un solo dataset y **sin "
         "jornada de holdout externa**: los 44 perfiles aparecen en las tres particiones",
         "Se recolecta jornada nueva el 24 de octubre. Hasta entonces **no se afirma "
         "generalización**"],
        ["**Validez de conclusión**", "Los intervalos por ventana son **descriptivos**: las "
         "ventanas de un episodio comparten historia y no son observaciones independientes",
         "Se reportan como descriptivos, no como prueba inferencial; la comparación entre "
         "modelos usa McNemar con corrección de Holm"],
        ["**Validez de constructo**", "*«¿el FPR de laboratorio mide el FPR de operación?»* — "
         "**la medición dice que no**: 4,71 % frente a 25,81 % (pase 1) y 22,97 % (pase 2)", "Se reportan ambos por "
         "separado y se declara la brecha como el hallazgo principal"],
    ], [3.4, 6.6, 7.4])

    # ---------------------------------------------------------------- 5
    h1(doc, "6 ·", "Conclusión")
    par(doc, "Se demostró con evidencia que el sistema **detecta y bloquea en tiempo real** sobre "
             "una red enrutada, con ROC-AUC de 0,974, detección del 88,8 % sobre ataques genuinos "
             "y bloqueo en una mediana de 8 segundos, sin ninguna caída de servicio registrada.")
    par(doc, "**No se demostró** que lo haga con una tasa de falsos positivos aceptable sobre "
             "tráfico legítimo pesado: en esa condición el sistema todavía no es apto para "
             "operación desatendida.")
    par(doc, "Delimitar esa frontera con medición —y no ocultarla— es el resultado de esta "
             "evaluación. Un hallazgo negativo verificado vale más que una conclusión favorable "
             "sin respaldo.", italic=True, color=DIM)
    bloque_enlaces(doc, "Evidencia en el repositorio", [
        ("Informe detallado, con las 11 figuras y la trazabilidad de cada cifra",
         "docs/entregables/01-evaluacion-critica/informe-evaluacion-critica.md"),
        ("Ablación por capas y comparación 14 vs. 28 variables",
         "docs/fase04-modelado/07-ablacion-multicapa.md"),
        ("Significancia estadística entre los siete modelos",
         "docs/fase04-modelado/08-significancia-entre-modelos.md"),
        ("Validación del sistema desplegado (F6)",
         "docs/fase07-validacion-final/02-resultados-f6.md"),
        ("Las 11 gráficas, generadas por script desde los datos reales",
         "docs/entregables/graficas"),
    ])

    rematar(doc,

            "Informe de evaluación crítica de resultados",

            "Investigación V · Sesión 01 · Evaluación crítica de los resultados de la tesis",

            "Informe de evaluación crítica · Salazar Tocas & Sauñe Fernandez",

            "Investigación V · Sesión 01 · UPeU")

    doc.save(OUT)
    print(f"Generado: {OUT.relative_to(REPO)}")


if __name__ == "__main__":
    main()
