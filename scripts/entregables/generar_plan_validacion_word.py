#!/usr/bin/env python3
"""Word del plan de validacion de resultados (Sesion 02, momento CREA).

Formato exigido: 1-2 paginas. El .md es la fuente versionable; este script
produce el documento presentable.

    python3 scripts/entregables/generar_plan_validacion_word.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "docs" / "entregables" / "assets" / "logo-upeu.png"
OUT = REPO / "docs/entregables/07-plan-de-validacion/Plan-de-validacion-de-resultados.docx"

INK = RGBColor(0x13, 0x1B, 0x2E)
DIM = RGBColor(0x5B, 0x6B, 0x8C)
ACCENT = RGBColor(0x0F, 0x8A, 0x7D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_HEAD, F_ZEBRA, F_OK = "0F8A7D", "F4F6FB", "E0F3E6"


def shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def h1(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(txt)
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    return p


def parrafo(doc, txt, size=9.2, italic=False, color=INK, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i, tramo in enumerate(txt.split("**")):
        r = p.add_run(tramo)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.italic = italic
        r.font.bold = i % 2 == 1
    return p


def tabla(doc, cabeceras, filas, anchos):
    t = doc.add_table(rows=1, cols=len(cabeceras))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c, (txt, w) in enumerate(zip(cabeceras, anchos)):
        cell = t.rows[0].cells[c]
        cell.width = Cm(w)
        cell.text = ""
        r = cell.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.size = Pt(8.4)
        r.font.color.rgb = WHITE
        shade(cell, F_HEAD)
    for i, fila in enumerate(filas):
        row = t.add_row()
        for c, txt in enumerate(fila):
            cell = row.cells[c]
            cell.width = Cm(anchos[c])
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            for k, tramo in enumerate(str(txt).split("**")):
                r = p.add_run(tramo)
                r.font.size = Pt(8.2)
                r.font.bold = k % 2 == 1
                r.font.color.rgb = INK
            if txt == "✔":
                shade(cell, F_OK)
            elif i % 2 == 0:
                shade(cell, F_ZEBRA)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t


def main() -> None:
    if not LOGO.exists():
        raise SystemExit(f"falta el logo: {LOGO}")
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(1.5)
    s.left_margin = s.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"

    doc.add_picture(str(LOGO), width=Cm(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sz, col, bold in [
        ("Universidad Peruana Unión", 10, INK, True),
        ("Facultad de Ingeniería y Arquitectura · E.P. de Ingeniería de Sistemas", 8.5, DIM, False),
        ("Investigación V · Sesión 02", 8.5, DIM, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.font.color.rgb = col
        r.font.bold = bold

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("PLAN DE VALIDACIÓN DE RESULTADOS")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Sistema open source para la detección temprana de comportamientos "
                  "anómalos en redes de datos\nRubén Mark Salazar Tocas · Uziel Elias Sauñe Fernandez")
    r.font.size = Pt(9)
    r.font.color.rgb = DIM

    parrafo(doc, "El producto es un sistema desplegado que **decide y bloquea tráfico en tiempo "
                 "real**, no un instrumento de medición por escalas. Por eso la confiabilidad se "
                 "valida sobre las decisiones del modelo y el determinismo del sistema: no existen "
                 "ítems de escala que correlacionar, así que **Alfa de Cronbach no aplica** al producto.")

    h1(doc, "1 · Confiabilidad")
    tabla(doc, ["Prueba", "Aplicada a", "Umbral aceptable"], [
        ["**Intervalo de Wilson 95 %**  ✔", "Toda proporción reportada",
         "Se reporta **siempre**; una familia con n < 10 no sostiene conclusión propia"],
        ["**McNemar exacto + Holm**  ✔", "Pares de modelos y de configuraciones de variables",
         "**p < 0,05 tras corregir.** Sin corrección, 21 comparaciones dan ≈66 % de probabilidad de un falso hallazgo"],
        ["**Validación cruzada agrupada por episodio**", "Modelo congelado OCSVM",
         "La detección media de los pliegues debe caer **dentro del intervalo de Wilson** de la evaluación de un solo paso"],
        ["**Bootstrap por episodio**, B = 1 000", "Umbral 1,8126087939765134",
         "**Coeficiente de variación < 5 %**; por encima se reporta como banda, no como valor puntual"],
        ["**Determinismo**", "Ajuste completo del modelo",
         "**SHA-256 idéntico** del .joblib en 10 ejecuciones"],
    ], [4.3, 4.5, 8.6])
    parrafo(doc, "Wilson se prefiere al intervalo normal con proporciones extremas y muestras "
                 "pequeñas —el caso de las familias con n = 6—. McNemar es la prueba pareada "
                 "correcta porque ambos modelos se evalúan sobre las mismas ventanas, y la binomial "
                 "exacta evita la aproximación ji² con recuentos bajos. Agrupar por episodio es "
                 "obligatorio: las ventanas del mismo episodio se solapan y repartirlas al azar "
                 "produciría fuga.", size=8.3, italic=True, color=DIM)

    h1(doc, "2 · Replicabilidad")
    tabla(doc, ["Elemento", "Dónde", ""], [
        ["Código", "GitHub, licencia **MIT**", "✔"],
        ["Dataset, manifiesto y **los 7 modelos candidatos**", "Mismo repositorio, licencia **CC BY 4.0**", "✔"],
        ["Integridad", "sha256sum -c docs/dataset/SHA256SUMS", "✔"],
        ["Entorno", "Versiones exactas de scikit-learn y numpy en el manifiesto congelado", "✔"],
        ["Documentación", "Datasheet de 11 secciones, model card, system card y diccionario de las 28 variables", "✔"],
        ["**Depósito citable con DOI**", "Zenodo, versionado y enlazado al repositorio", "pendiente"],
        ["**Semillas y determinismo como protocolo**", "Protocolo de modelado", "pendiente"],
    ], [5.6, 9.4, 2.4])
    parrafo(doc, "**Umbral aceptable:** quien clone el repositorio debe reproducir el umbral "
                 "1,8126087939765134 **en sus 16 dígitos** y los recuentos 13/276 de falso positivo "
                 "y 158/179 de detección. Cualquier desviación invalida la replicación. El criterio "
                 "ya se cumple internamente: los siete modelos publicados, recargados y verificados "
                 "por SHA-256, reproducen el manifiesto de forma exacta.")

    h1(doc, "3 · Pertinencia")
    parrafo(doc, "Es el **único eje sin evidencia** a la fecha.", after=3)
    tabla(doc, ["Método", "Instrumento", "Umbral aceptable"], [
        ["**Prueba de usabilidad**", "System Usability Scale, 10 ítems, **5–8 evaluadores** con perfil de administración de redes",
         "**SUS ≥ 68**, media de referencia de la literatura. Por debajo, se rediseña el panel antes de la defensa"],
        ["**Tareas observadas**", "4 tareas: localizar una IP bloqueada, leer su expiración, distinguir alerta del modelo de alerta heurística, verificar servicios",
         "**Tasa de éxito ≥ 80 %** sin ayuda del observador"],
        ["**Entrevista semiestructurada**", "2–3 interesados: asesores y responsable de red",
         "Necesidades cubiertas y no cubiertas; evidencia cualitativa, sin umbral numérico"],
        ["**Trazabilidad de requisitos**", "Matriz de requisitos del jurado frente a la solución",
         "**100 % de filas cerradas**: con evidencia o declaradas fuera de alcance"],
    ], [3.6, 6.8, 7.0])
    parrafo(doc, "Se usa SUS y no una encuesta propia porque es un instrumento validado, con baremo "
                 "publicado, aplicable con muestras pequeñas y comparable con otros estudios.",
            size=8.3, italic=True, color=DIM)

    h1(doc, "Cronograma")
    tabla(doc, ["Sem.", "Eje", "Actividad", "Producto verificable"], [
        ["1", "Confiabilidad", "Validación cruzada por episodio y bootstrap del umbral", "Media y desviación por pliegue; banda del umbral"],
        ["1", "Replicabilidad", "Declarar semillas y determinismo", "Sección en el protocolo de modelado"],
        ["2", "Pertinencia", "Diseñar instrumento y guion de tareas; reclutar evaluadores", "Instrumento SUS y guion de sesión"],
        ["2", "Replicabilidad", "Depósito en Zenodo", "DOI citable"],
        ["3", "Pertinencia", "Ejecutar sesiones de usabilidad y entrevistas", "Puntaje SUS, tiempos y tasa de éxito"],
        ["3", "Pertinencia", "Cerrar la matriz de trazabilidad", "Matriz sin filas abiertas"],
        ["4", "Los tres", "Integrar la validación al artículo y a la tesis", "Sección de validación y limitaciones"],
    ], [1.3, 2.8, 6.6, 6.7])
    parrafo(doc, "El orden responde a dependencias: confiabilidad va primero porque no necesita a "
                 "nadie más; pertinencia ocupa dos semanas porque depende de coordinar personas.",
            size=8.3, italic=True, color=DIM)

    h1(doc, "Fuera del alcance de este plan")
    parrafo(doc, "El **falso positivo operativo de 23–26 %** sobre tráfico legítimo pesado exige "
                 "recalibrar con ese tráfico como normalidad y repetir la validación: semanas. La "
                 "**selección posterior del modelo** exige un protocolo nuevo con evaluación "
                 "reservada. Ambas se declaran como límite, no se simulan.")
    parrafo(doc, "✔ = ya ejecutado y publicado.", size=8, italic=True, color=DIM)

    doc.save(OUT)
    print(f"Generado: {OUT.relative_to(REPO)}")


if __name__ == "__main__":
    main()
