#!/usr/bin/env python3
"""Piezas de formato compartidas por los generadores de Word.

Existe para que los cuatro entregables se vean como una misma familia y para
no repetir en cada script el pie de pagina, los bordes de tabla y las
propiedades del documento.
"""
from __future__ import annotations
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

GRIS_LINEA = "BFBFBF"

# ---------------------------------------------------------------- PALETA ----
# Tema AZUL, comun a los entregables del curso. El acento es la identidad del
# documento; los rellenos de semaforo son SEMANTICOS y por eso no se tinen de
# azul: verde/ambar/rojo comunican estado, no tema.
AZUL = "1F4E79"          # acento: titulos, cabeceras de tabla, bordes de caja
AZUL_CLARO = "2E74B5"    # enlaces y realces secundarios
F_HEAD = AZUL            # relleno de cabecera de tabla
F_ZEBRA = "EEF3FA"       # filas alternas, tinte azul muy claro
F_CAJA = "F2F6FC"        # fondo de caja destacada
F_OK, F_AMBER, F_DANGER = "E0F3E6", "FDECD2", "FBE3E1"   # semaforo, semantico

DIM = RGBColor(0x5B, 0x6B, 0x8C)
INK = RGBColor(0x13, 0x1B, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_CLARO = RGBColor(0x2E, 0x74, 0xB5)

REPO_URL = ("https://github.com/marksato13/"
            "VF-Sistema-Open-Source-para-la-Deteccion-Temprana-de-Comportamientos-Anomalos-en-Redes-de-Datos")


def url_repo(ruta: str = "") -> str:
    """URL navegable a un archivo o carpeta del repositorio."""
    return f"{REPO_URL}/blob/main/{ruta}" if ruta else REPO_URL


def enlace(parrafo, texto: str, url: str, size: float = 8.6) -> None:
    """Hipervinculo real: python-docx no lo soporta de fabrica."""
    rid = parrafo.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), AZUL_CLARO)
    und = OxmlElement("w:u"); und.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
    for e in (col, und, sz):
        rPr.append(e)
    t = OxmlElement("w:t"); t.text = texto
    r.append(rPr); r.append(t); h.append(r)
    parrafo._p.append(h)


def bordes_tabla(tabla, color: str = GRIS_LINEA, grosor: int = 4) -> None:
    """Retícula fina. Sin bordes, las celdas sombreadas parecen flotar."""
    tblPr = tabla._tbl.tblPr
    for viejo in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(viejo)
    b = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{lado}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(grosor))
        e.set(qn("w:color"), color)
        b.append(e)
    tblPr.append(b)


def _campo(parrafo, instruccion: str, size: float, color: RGBColor):
    """Inserta un campo de Word (PAGE, NUMPAGES) que se recalcula al abrir."""
    r = parrafo.add_run()
    r.font.size = Pt(size)
    r.font.color.rgb = color
    ini = OxmlElement("w:fldChar")
    ini.set(qn("w:fldCharType"), "begin")
    txt = OxmlElement("w:instrText")
    txt.set(qn("xml:space"), "preserve")
    txt.text = f" {instruccion} "
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for el in (ini, txt, fin):
        r._r.append(el)


def pie_con_pagina(doc, izquierda: str, size: float = 7.4) -> None:
    """Pie con identificacion a la izquierda y 'Pagina X de Y' a la derecha."""
    p = doc.sections[0].footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # tabulador derecho al ancho util de la pagina
    s = doc.sections[0]
    ancho_emu = int(s.page_width) - int(s.left_margin) - int(s.right_margin)
    ancho_twips = ancho_emu // 635  # 1 twip = 635 EMU
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(ancho_twips))
    tabs.append(tab)
    pPr.append(tabs)

    r = p.add_run(izquierda + "\t")
    r.font.size = Pt(size)
    r.font.color.rgb = DIM
    r = p.add_run("Página ")
    r.font.size = Pt(size)
    r.font.color.rgb = DIM
    _campo(p, "PAGE", size, DIM)
    r = p.add_run(" de ")
    r.font.size = Pt(size)
    r.font.color.rgb = DIM
    _campo(p, "NUMPAGES", size, DIM)


def encabezado(doc, texto: str, size: float = 7.4) -> None:
    """Encabezado discreto, con una linea inferior de separacion."""
    p = doc.sections[0].header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.font.color.rgb = DIM
    r.font.italic = True
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "4")
    e.set(qn("w:color"), GRIS_LINEA)
    e.set(qn("w:space"), "3")
    bd.append(e)
    pPr.append(bd)


def propiedades(doc, titulo: str, asunto: str,
                autor: str = "Rubén Mark Salazar Tocas; Uziel Elias Sauñe Fernandez",
                palabras: str = "detección de anomalías; seguridad de redes; "
                                "aprendizaje no supervisado; OCSVM") -> None:
    """Metadatos del archivo. Un .docx sin titulo ni autor se ve improvisado
    en cuanto alguien abre las propiedades o lo indexa un gestor documental."""
    cp = doc.core_properties
    cp.title = titulo
    cp.subject = asunto
    cp.author = autor
    cp.last_modified_by = autor
    cp.keywords = palabras
    cp.category = "Investigación V · UPeU"
    cp.comments = ("Documento generado por script desde los artefactos del proyecto; "
                   "ninguna cifra se transcribe a mano.")


def bloque_enlaces(doc, titulo: str, entradas: list[tuple[str, str]]) -> None:
    """Bloque final con hipervinculos navegables al repositorio.

    Un documento que cita "ver el repositorio" sin dar la URL obliga al lector
    a buscarla. Aqui cada referencia es un enlace que se abre con un clic.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(titulo)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    for etiqueta, ruta in entradas:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(1)
        q.paragraph_format.left_indent = Pt(10)
        a = q.add_run(f"{etiqueta} — ")
        a.font.size = Pt(8.4)
        a.font.color.rgb = INK
        enlace(q, "abrir en GitHub", url_repo(ruta), size=8.4)


def rematar(doc, titulo: str, asunto: str, pie_izquierda: str, encabezado_txt: str) -> None:
    """Aplica de una vez todo el mobiliario del documento."""
    propiedades(doc, titulo, asunto)
    pie_con_pagina(doc, pie_izquierda)
    encabezado(doc, encabezado_txt)
    for t in doc.tables:
        bordes_tabla(t)
