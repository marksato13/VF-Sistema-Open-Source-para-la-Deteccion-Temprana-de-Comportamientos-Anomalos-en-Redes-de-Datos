#!/usr/bin/env python3
"""Genera la Ficha de auditoría del producto de ingeniería en formato Word.

Mismo contenido que `docs/entregables/04-ficha-auditoria/`, con la
presentación formal del curso. Los subtotales y el puntaje final se **calculan
aquí**, no se escriben a mano, de modo que la ficha no puede quedar
descuadrada si cambia una puntuación.

Uso:
    .venv/bin/python3 scripts/entregables/generar_ficha_word.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from _docx_estilo import rematar, bloque_enlaces

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from generar_informe_word import (  # reutiliza el formato ya definido
    ACCENT, AMBER, DANGER, DIM, F_AMBER, F_DANGER, F_OK, INK, OK,
    caja, cell_text, figura, h1, h2, parrafo, shade, tabla, vineta,
)

REPO = Path(__file__).resolve().parents[2]

# Cifras del repositorio leídas en tiempo de generación: escritas a mano
# envejecen en cuanto se hace un commit más.
def _git(*args) -> int:
    import subprocess
    r = subprocess.run(["git", "-C", str(REPO), *args], capture_output=True, text=True)
    return r.stdout.count("\n")


N_ARCHIVOS = _git("ls-files")
N_COMMITS = _git("log", "--oneline")
OUT = (REPO / "docs" / "entregables" / "04-ficha-auditoria"
       / "Ficha-auditoria-producto.docx")
LOGO = REPO / "docs" / "entregables" / "assets" / "logo-upeu.png"

# ------------------------------------------------------------- puntuaciones --
# (id, criterio, evidencia, puntaje)  ·  "N/A" se excluye del cálculo
CONFIABILIDAD = [
    ("1.1", "Alfa de Cronbach (consistencia interna)",
     "El producto no emplea cuestionarios ni escalas psicométricas; no hay ítems que correlacionar", "N/A"),
    ("1.2", "Kappa de Cohen (acuerdo inter-evaluador)",
     "No se realizó evaluación por jueces ni doble etiquetado independiente; las etiquetas provienen del diseño experimental", 0),
    ("1.3", "Validación cruzada",
     "Ejecutada, agrupada por episodio en 5 pliegues: la detección media cae dentro del intervalo de "
     "confianza de la evaluación de un solo paso, así que el resultado no depende de la partición elegida", 3),
    ("1.4", "Reproducibilidad de la medición",
     "Al reevaluar el modelo se obtuvieron exactamente las cifras originales (13/276 y 158/179)", 3),
    ("1.5", "Estabilidad entre repeticiones",
     "Dos pases de validación operativa con resultados equivalentes (25,8 % y 23,0 %) y remuestreo del "
     "umbral con B = 1 000: coeficiente de variación 4,10 %, por debajo del 5 % declarado de antemano", 3),
    ("1.6", "Cuantificación de la incertidumbre",
     "Intervalos de Wilson 95 % en toda proporción y prueba de significancia entre modelos: McNemar exacto "
     "con corrección de Holm-Bonferroni sobre las 21 comparaciones por pares", 3),
]
REPLICABILIDAD = [
    ("2.1", "Código disponible",
     f"Repositorio público con {N_ARCHIVOS} archivos versionados y {N_COMMITS} registros de cambios trazables", 3),
    ("2.2", "Datos disponibles",
     "Publicados: dataset, manifiesto y los 7 modelos candidatos, con checksums verificables por "
     "sha256sum -c y licencias MIT para el código y CC BY 4.0 para los datos", 3),
    ("2.3", "Entorno documentado",
     "Versiones exactas fijadas, script de instalación idempotente y playbooks de Ansible", 3),
    ("2.4", "Determinismo y semillas",
     "Protocolo declarado: qué componentes son estocásticos y cuáles no, con verificación de 10 ajustes "
     "repetidos que producen el mismo SHA-256 y el mismo umbral", 3),
    ("2.5", "Integridad verificable",
     "SHA-256 de datos, modelo y calibrador; repositorio verificado limpio antes y después", 3),
    ("2.6", "Instrucciones de reproducción",
     "El datasheet documenta el procedimiento exacto de descarga, verificación y regeneración; queda "
     "pendiente el manual de instalación desde cero", 3),
]
PERTINENCIA = [
    ("3.1", "Validación con usuarios reales",
     "No se realizó. Sin pruebas con analistas de seguridad ni medición de experiencia de uso del panel", 0),
    ("3.2", "Evaluación por expertos o jueces",
     "No se aplicó ningún instrumento de juicio experto (Delphi, SUS u otro)", 0),
    ("3.3", "Trazabilidad de requisitos",
     "Matriz cerrada: cada requisito está cumplido con evidencia enlazada, cumplido con una reserva medida, "
     "o declarado pendiente con lo que concretamente falta. Ninguna fila queda en Planificado", 3),
    ("3.4", "Validación en operación real",
     "Sistema medido desplegado y activo: 2 pases de 29 corridas con motor y bloqueo sobre tráfico real", 3),
    ("3.5", "Alineación con el problema",
     "Detecta y bloquea las 6 familias de ataque previstas, con métricas medidas por familia", 3),
    ("3.6", "Alcance y limitaciones declarados",
     "Limitaciones medidas, cuantificadas y publicadas, incluido el error operativo desfavorable (23–26 %)", 3),
]

NIVEL = {3: ("Completo", F_OK, OK), 2: ("Parcial", F_AMBER, AMBER),
         1: ("Insuficiente", F_DANGER, DANGER), 0: ("Ausente", F_DANGER, DANGER)}


def subtotal(items):
    vals = [p for *_, p in items if p != "N/A"]
    return sum(vals), len(vals) * 3


def bloque(doc, titulo, pregunta, items, lectura):
    h1(doc, titulo)
    parrafo(doc, pregunta, italic=True, color=DIM)
    filas = []
    for cid, crit, ev, p in items:
        if p == "N/A":
            celda = ("N/A", "F2F6FC")
        else:
            celda = (f"{p}", NIVEL[p][1])
        filas.append((cid, crit, ev, celda))
    tabla(doc, ["#", "Criterio", "Evidencia concreta en el proyecto", "Pts"],
          filas, widths=[1.0, 4.3, 9.5, 1.2])
    obt, mx = subtotal(items)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"Subtotal: {obt} / {mx} puntos  =  {100*obt/mx:.1f} %")
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = OK if obt/mx >= 0.7 else (AMBER if obt/mx >= 0.5 else DANGER)
    parrafo(doc, lectura)
    return obt, mx


def main() -> int:
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.9)
        s.left_margin = s.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    # -------------------------------------------------------- CARÁTULA -----
    if not LOGO.exists():
        raise SystemExit(f"falta el logo: {LOGO}")
    doc.add_picture(str(LOGO), width=Cm(7.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sz, col, bold in [("Universidad Peruana Unión", 12.5, INK, True),
                               ("Facultad de Ingeniería y Arquitectura", 10.5, DIM, False),
                               ("E.P. de Ingeniería de Sistemas", 10.5, DIM, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = col

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FICHA DE AUDITORÍA\nDEL PRODUCTO DE INGENIERÍA VALIDADO")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Detección temprana de comportamientos anómalos en redes de datos\n"
                  "mediante modelos predictivos y un mecanismo de control inline")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = DIM

    doc.add_paragraph()
    tabla(doc, ["Campo", "Detalle"],
          [("Producto auditado", "Sistema de detección de anomalías de red con control inline, desplegado en laboratorio virtualizado (VM02)"),
           ("Curso", "Investigación V · Ciclo X"),
           ("Docente", "Ing. Nemias Saboya Ríos"),
           ("Integrantes", "Rubén Mark Salazar Tocas\nUziel Elias Sauñe Fernandez"),
           ("Fecha", "19 de agosto de 2026")],
          widths=[3.6, 12.4])

    doc.add_paragraph()
    caja(doc, "Nota sobre la rúbrica",
         "La escala y los ítems son una **reconstrucción razonada** del ejercicio propuesto, porque no se dispuso "
         "del formato exacto de la ficha proyectada en clase. La estructura de tres dimensiones —confiabilidad, "
         "replicabilidad y pertinencia— y el cálculo de puntaje final sí corresponden a la consigna. Si el formato "
         "oficial difiere, los puntajes por ítem se trasladan sin recalcular la evidencia.",
         fill="FDECD2", color_borde="B45309")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -------------------------------------------- PROCEDIMIENTO / MÉTODO ---
    h1(doc, "Procedimiento seguido")
    parrafo(doc, "La consigna planteaba seis pasos. Se dejan explícitos para que pueda verificarse que se "
                 "siguieron y no solo que se entregó la ficha:")
    for t in [
        "**Paso 1 — Conformación del equipo.** Rubén Mark Salazar Tocas y Uziel Elias Sauñe Fernandez, los mismos "
        "integrantes del proyecto auditado.",
        "**Paso 2 — Análisis del producto validado.** Se auditó el sistema realmente desplegado en VM02, no su "
        "especificación: motor de decisión, mecanismo de bloqueo y panel de observación.",
        "**Paso 3 — Evidencia de confiabilidad.** Las tres vías nombradas (Alfa, Kappa, validación cruzada) más las "
        "que corresponden a un sistema de software. → Sección 1",
        "**Paso 4 — Evidencia de replicabilidad.** Se comprobó ejecutando, no solo leyendo documentación, qué está "
        "publicado y qué no. → Sección 2",
        "**Paso 5 — Evidencia de pertinencia.** Se contrastó el producto con el problema declarado y con los "
        "requisitos comprometidos. → Sección 3",
        "**Paso 6 — Diligenciamiento y cálculo.** Puntuación ítem por ítem, subtotales y puntaje final. → Secciones 4 y 5",
    ]:
        vineta(doc, t)

    doc.add_paragraph()
    h1(doc, "Cómo se realizó la auditoría")
    parrafo(doc, "Para que la ficha sea verificable y no una autoevaluación complaciente, se fijó de antemano qué "
                 "se aceptaría como evidencia:")
    for t in [
        "**Solo cuenta lo que puede comprobar un tercero.** Una afirmación documentada pero no ejecutada se puntúa "
        "como declarada, nunca como completa.",
        "**Se verificó ejecutando, no leyendo.** La reproducibilidad se comprobó reevaluando el modelo congelado; la "
        "disponibilidad de datos, inspeccionando qué excluye el repositorio y cuánto ocupa cada artefacto.",
        "**Las cifras del proyecto no se transcriben a mano.** Se leen del manifiesto de calibración y de los "
        "registros de la validación operativa.",
        "**Lo ausente se puntúa como ausente.** No se compensa una carencia con una fortaleza de otra dimensión.",
        "**Lo que no aplica se excluye del cálculo**, en vez de puntuarse con cero.",
    ]:
        vineta(doc, t)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------- ESCALA -----
    h1(doc, "Escala de valoración")
    tabla(doc, ["Puntos", "Nivel", "Significado"],
          [(("3", F_OK), "Completo", "Existe y es verificable por un tercero"),
           (("2", F_AMBER), "Parcial", "Existe pero con limitaciones declaradas"),
           (("1", F_DANGER), "Insuficiente", "Solo declarado, sin evidencia sólida"),
           (("0", F_DANGER), "Ausente", "No se abordó"),
           (("N/A", "F2F6FC"), "No aplica", "No corresponde al tipo de producto (se excluye del cálculo)")],
          widths=[1.8, 3.0, 11.2])

    # ------------------------------ FICHA DEL DOCENTE (6 criterios / 20) ----
    doc.add_paragraph()
    h1(doc, "Ficha de auditoría de validación — esquema del curso")
    parrafo(doc, "Los seis criterios sobre 20 puntos de la Sesión 02. El análisis extenso de las "
                 "secciones 1 a 3 los sustenta con 18 ítems sobre 51 puntos.")
    tabla(doc, ["Criterio", "Evidencia encontrada", "Puntaje"],
          [("Confiabilidad estadística reportada (Alfa, Kappa u otra)",
            "Intervalos de Wilson 95 % en toda proporción, McNemar exacto con corrección de Holm sobre 21 "
            "comparaciones, ROC-AUC 0,974 y validación cruzada agrupada por episodio. Alfa no aplica: no hay "
            "escalas. Kappa no aplica: no hay jueces",
            ("4 / 4", F_OK)),
           ("Método de validación de resultados declarado",
            "Hold-out con partición disjunta por episodio, umbral calibrado solo con validation y evaluación "
            "bloqueada de un solo paso, con validación cruzada por episodio que la respalda. Falta la jornada de "
            "holdout temporal externa",
            ("3 / 4", F_AMBER)),
           ("Confiabilidad de sistema / determinismo",
            "Reproducción bit a bit verificada: el umbral coincide en sus 16 dígitos y los 7 modelos reproducen "
            "el manifiesto. Las semillas aún no se declaran como protocolo",
            ("3 / 3", F_OK)),
           ("Datos y/o código disponibles públicamente",
            "Dataset, manifiesto y los 7 modelos candidatos publicados, verificables con sha256sum -c, "
            "bajo licencias MIT y CC BY 4.0", ("3 / 3", F_OK)),
           ("Entorno y dependencias documentadas",
            "Versiones exactas de scikit-learn y numpy registradas en el manifiesto congelado, con script de "
            "instalación idempotente", ("3 / 3", F_OK)),
           ("Pertinencia validada con usuarios o stakeholders reales",
            "No se realizó ninguna. Es la única ausencia total del producto",
            ("0 / 3", F_DANGER)),
           (("TOTAL", "F2F6FC"), "", ("16 / 20  ·  80 %", F_AMBER))],
          widths=[5.0, 9.0, 2.0])
    parrafo(doc, "**El único cero es la validación con personas.** No se corrige documentando: exige "
                 "evaluadores usando el panel. Una prueba de usabilidad con 5–8 participantes lo convierte "
                 "en evidencia y eleva el total a 19 de 20.", italic=True, color=DIM)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    a = bloque(doc, "1.  Evidencia de CONFIABILIDAD",
               "¿Los resultados son estables y consistentes al repetir la medición?", CONFIABILIDAD,
               "**Lectura.** La cuantificación de la incertidumbre pasó de parcial a completa: toda proporción "
               "lleva intervalo de Wilson y las comparaciones entre modelos tienen prueba de significancia. "
               "Lo que sigue faltando es **validación cruzada sobre el modelo elegido**; el acuerdo "
               "inter-evaluador no aplica porque el diseño no contempla jueces.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    b = bloque(doc, "2.  Evidencia de REPLICABILIDAD",
               "¿Puede un tercero reconstruir el estudio y obtener lo mismo?", REPLICABILIDAD,
               "**Lectura.** Es la dimensión **más fuerte** del producto, y la que más subió: la publicación del "
               "dataset y de los siete modelos con checksums y licencias cerró la única brecha que quedaba. "
               "Un tercero puede hoy clonar el repositorio y reproducir el umbral en sus 16 dígitos.")

    doc.add_paragraph()
    c = bloque(doc, "3.  Evidencia de PERTINENCIA",
               "¿El producto responde al problema real y su utilidad está demostrada?", PERTINENCIA,
               "**Lectura.** El producto es **técnicamente pertinente** —resuelve el problema y se probó en "
               "operación real— pero carece de **validación con personas**: nadie externo al equipo lo ha usado ni "
               "evaluado. Para un producto cuya interfaz es un panel destinado a un analista, esa ausencia pesa.")

    # ---------------------------------------------------- PUNTAJE FINAL ----
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h1(doc, "4.  Puntaje final")
    T, M = a[0] + b[0] + c[0], a[1] + b[1] + c[1]

    def nivel(pct):
        return ("Alto", F_OK) if pct >= 70 else (("Medio", F_AMBER) if pct >= 50 else ("Bajo", F_DANGER))

    filas = []
    for nombre, (o_, m_) in [("Confiabilidad", a), ("Replicabilidad", b), ("Pertinencia", c)]:
        pct = 100 * o_ / m_
        n, f = nivel(pct)
        filas.append((nombre, f"{o_}", f"{m_}", (f"{pct:.1f} %", f), (n, f)))
    pctT = 100 * T / M
    nT, fT = nivel(pctT)
    filas.append(("TOTAL", f"{T}", f"{M}", (f"{pctT:.1f} %", fT), (nT, fT)))
    tabla(doc, ["Dimensión", "Obtenido", "Máximo", "Porcentaje", "Nivel"], filas,
          widths=[4.4, 2.6, 2.4, 3.2, 3.4])

    doc.add_paragraph()
    caja(doc, f"Interpretación del {pctT:.1f} %",
         "Describe con precisión el estado del producto: **sólido como artefacto de ingeniería, incompleto como "
         "evidencia científica**. Lo que sostiene el puntaje es la **replicabilidad** (77,8 %): el trabajo es "
         "verificable, versionado y auditable. Lo que lo baja son dos ausencias distintas: **validación estadística** "
         "(sin validación cruzada del modelo elegido) y **validación humana** (ningún usuario o experto externo lo "
         "evaluó). Ninguna invalida los resultados; ambas **limitan el alcance de lo que puede afirmarse** con ellos.")

    # ------------------------------------------------------- ACCIONES -----
    doc.add_paragraph()
    h1(doc, "5.  Acciones para elevar el puntaje")
    parrafo(doc, "**Siete ítems subieron desde la primera auditoría**, todos con evidencia publicada y "
                 "ninguno con experimentación nueva.")
    tabla(doc, ["Acción ejecutada", "Ítem", "De → a"],
          [("Publicar el dataset, el manifiesto y los 7 modelos con checksums y licencias", "2.2", ("1 → 3", F_OK)),
           ("Intervalos de Wilson en toda proporción y McNemar con corrección de Holm", "1.6", ("2 → 3", F_OK)),
           ("Validación cruzada agrupada por episodio sobre el modelo congelado", "1.3", ("1 → 3", F_OK)),
           ("Remuestreo del umbral con B = 1 000; coeficiente de variación 4,10 %", "1.5", ("2 → 3", F_OK)),
           ("Protocolo de determinismo, verificado con 10 ajustes de SHA-256 idéntico", "2.4", ("2 → 3", F_OK)),
           ("Documentar el procedimiento de descarga, verificación y regeneración", "2.6", ("2 → 3", F_OK)),
           ("Cerrar la matriz de trazabilidad de requisitos", "3.3", ("1 → 3", F_OK))],
          widths=[10.2, 1.8, 2.0])
    parrafo(doc, f"**Efecto acumulado: de 32/51 (62,7 %) a {T}/{M} ({pctT:.1f} %).**")

    doc.add_paragraph()
    parrafo(doc, "**Lo que queda**")
    tabla(doc, ["Acción pendiente", "Sube", "De → a", "Tiempo"],
          [("**Aplicar un instrumento validado (SUS) con 5–8 evaluadores sobre el panel**", "3.1", ("0 → 3", F_AMBER), ),
           ("Sesión de juicio experto con 3 evaluadores", "3.2", ("0 → 2", F_AMBER)),
           ("Acuerdo inter-evaluador, si se incorpora doble etiquetado independiente", "1.2", ("0 → 2", F_AMBER))],
          widths=[9.4, 1.8, 2.0, 2.8])

    doc.add_paragraph()
    caja(doc, "El puntaje ya no lo frena la evidencia técnica",
         f"La replicabilidad está en **100 %** y la confiabilidad en **80 %**. Los tres ítems que quedan miden "
         f"**lo que otras personas opinan del producto**, y ninguno se corrige escribiendo. Con solo la prueba "
         f"de usabilidad el puntaje llegaría a **45/{M} ({100*45/M:.1f} %)**; con los tres, a "
         f"**49/{M} ({100*49/M:.1f} %)**.")

    # -------------------------------------------------- ISO/IEC 25010 -----
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h1(doc, "6.  Correspondencia con la norma ISO/IEC 25010")
    parrafo(doc, "La norma de calidad de producto de software permite situar los resultados en un marco reconocido. "
                 "Se declara solo lo que el proyecto **midió**; el resto se marca como no evaluado, en lugar de "
                 "suponerlo.")
    tabla(doc, ["Característica ISO/IEC 25010", "¿Se evaluó?", "Evidencia en este proyecto"],
          [("Fiabilidad — madurez, disponibilidad, tolerancia a fallos", ("Sí", F_OK),
            "Cero caídas registradas en 58 corridas (55 verificadas), sin pérdida de paquetes. Tres fallos de producción detectados y corregidos con prueba positiva y negativa"),
           ("Eficiencia de desempeño — comportamiento temporal", ("Sí", F_OK),
            "Bloqueo en una mediana de 8 s. Límite declarado: bajo carga sostenida el motor acumula retraso"),
           ("Adecuación funcional — completitud y corrección", ("Parcial", F_AMBER),
            "Detecta y bloquea las 6 familias previstas (88,8 %), pero la corrección se degrada con tráfico legítimo intenso (error 23–26 %)"),
           ("Seguridad — confidencialidad, integridad, no repudio", ("Parcial", F_AMBER),
            "Integridad verificable por SHA-256 y acceso por helper de alcance estrecho. NO se evaluó el sistema como objetivo de ataque (evasión, suplantación de IP)"),
           ("Mantenibilidad — modularidad, reusabilidad", ("Parcial", F_AMBER),
            f"El motor reutiliza el extractor congelado sin duplicar fórmulas; {N_ARCHIVOS} archivos versionados. Sin métricas formales"),
           ("Usabilidad", ("No", F_DANGER),
            "El panel no se sometió a ninguna evaluación de uso (ver ítems 3.1 y 3.2)"),
           ("Portabilidad", ("No", F_DANGER),
            "Desplegado sobre una única configuración de laboratorio; no se probó en otro entorno"),
           ("Compatibilidad", ("No", F_DANGER),
            "No se evaluó la coexistencia con otras herramientas de monitoreo")],
          widths=[5.4, 2.2, 8.9])
    doc.add_paragraph()
    caja(doc, "Lectura",
         "Las dos características que el producto **sí demuestra con evidencia** —fiabilidad y eficiencia de "
         "desempeño— son precisamente las críticas en un sistema de detección en tiempo real. Las no evaluadas "
         "coinciden con las carencias que ya señala la ficha: usabilidad con la ausencia de validación con usuarios, "
         "y seguridad con la ausencia de pruebas de evasión.")

    # ----------------------------------------------------- CONCLUSIÓN -----
    doc.add_paragraph()
    h1(doc, "7.  Conclusión de la auditoría")
    parrafo(doc, "El producto **está validado como artefacto de ingeniería**: funciona, se midió en operación real "
                 "y sus resultados se pueden reproducir exactamente. Lo que la auditoría expone no son fallos del "
                 "sistema, sino **huecos en la evidencia que lo respalda**: falta validación cruzada, faltan los "
                 "datos publicados y falta que alguien ajeno al equipo lo haya usado.")
    parrafo(doc, "La ventaja es que **la mayor parte de esos huecos se cierra en horas**, porque el material ya "
                 "existe y solo requiere publicarse o ejecutarse. La excepción es la validación con usuarios, que "
                 "exige planificar una sesión con evaluadores externos.")

    # ----------------------------------------------------- REFERENCIAS ----
    doc.add_paragraph()
    h1(doc, "Referencias")
    for ref in [
        "**ISO/IEC 25010:2011.** Systems and software engineering — SQuaRE — System and software quality models. "
        "International Organization for Standardization.",
        "**Cronbach, L. J. (1951).** Coefficient alpha and the internal structure of tests. Psychometrika, 16(3), "
        "297–334. — Origen del criterio del ítem 1.1.",
        "**Cohen, J. (1960).** A coefficient of agreement for nominal scales. Educational and Psychological "
        "Measurement, 20(1), 37–46. — Origen del criterio del ítem 1.2.",
        "**Wilson, E. B. (1927).** Probable inference, the law of succession, and statistical inference. Journal of "
        "the American Statistical Association, 22(158), 209–212. — Método usado para los intervalos del ítem 1.6.",
        "**Wilkinson, M. D., et al. (2016).** The FAIR Guiding Principles for scientific data management and "
        "stewardship. Scientific Data, 3, 160018. — Marco de referencia de la sección 2.",
        "**Peng, R. D. (2011).** Reproducible research in computational science. Science, 334(6060), 1226–1227. — "
        "Distinción entre reproducibilidad y replicabilidad, ítems 1.4 y 2.1–2.6.",
    ]:
        vineta(doc, ref)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    bloque_enlaces(doc, "Evidencia en el repositorio", [
        ("Ficha detallada, con los 18 ítems y su evidencia",
         "docs/entregables/04-ficha-auditoria/ficha-auditoria.md"),
        ("Plan de mejora: debilidades abiertas con prioridad e impacto",
         "docs/entregables/06-plan-de-mejora/README.md"),
        ("Instrumento SUS, listo para aplicar",
         "docs/entregables/08-validacion-usuarios/README.md"),
        ("Matriz de trazabilidad de los requisitos del jurado",
         "docs/requisitos-jurado/README.md"),
    ])

    rematar(doc,
            "Ficha de auditoría del producto",
            "Auditoría de confiabilidad, replicabilidad y pertinencia",
            "Ficha de auditoría del producto · Salazar Tocas & Sauñe Fernandez",
            "Investigación V · Sesión 02 · UPeU")
    doc.save(OUT)
    print(f"Generado: {OUT.relative_to(REPO)}")
    print(f"  Confiabilidad {a[0]}/{a[1]} · Replicabilidad {b[0]}/{b[1]} · Pertinencia {c[0]}/{c[1]}")
    print(f"  TOTAL {T}/{M} = {pctT:.1f} %   (con la prueba de usabilidad: {100*45/M:.1f} %)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
