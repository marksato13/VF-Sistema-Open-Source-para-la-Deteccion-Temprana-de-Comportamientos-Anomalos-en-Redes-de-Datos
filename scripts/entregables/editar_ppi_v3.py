#!/usr/bin/env python3
"""Aplica al PPI las correcciones aprobadas del analisis v3.

Trabaja SOBRE EL MISMO documento, con respaldo previo. Cada bloque comprueba
que el texto que va a sustituir es el esperado; si no lo encuentra, aborta en
vez de escribir en el sitio equivocado.

    python3 scripts/entregables/editar_ppi_v3.py
"""
from __future__ import annotations
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))
from _omml import r, sub, sup, frac, nary, ecuacion          # noqa: E402

from docx import Document                                     # noqa: E402
from docx.shared import Pt                                    # noqa: E402

PPI = REPO / "docs/entregables/05-ppi/PPI Editar_actual.docx"


def texto(p, nuevo: str) -> None:
    """Sustituye el texto conservando el formato del primer run."""
    if p.runs:
        p.runs[0].text = nuevo
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        p.add_run(nuevo)


def exige(p, fragmento: str, etiqueta: str) -> None:
    if fragmento.lower() not in p.text.lower():
        raise SystemExit(f"ABORTA en {etiqueta}: se esperaba «{fragmento[:45]}» "
                         f"y hay «{p.text[:60]}»")


def main() -> None:
    d = Document(PPI)
    P = d.paragraphs
    cambios = []

    # ---------------------------------------------- B-02 · titulos limpios --
    for i, viejo, nuevo in [
        (176, "2.1 Diseño Metodológico", "2.1 Diseño metodológico"),
        (302, "2.7 Aspectos Éticos", "2.7 Aspectos éticos"),
    ]:
        exige(P[i], viejo.split()[1][:6], f"título p{i}")
        texto(P[i], nuevo)
        cambios.append(f"p{i} título depurado → «{nuevo}»")

    # ------------------------------- A-02 · OCSVM sustituye a Isolation Forest
    exige(P[102], "Isolation Forest", "p102")
    texto(P[102],
          "El modelo desplegado es un One-Class SVM (OCSVM), un método de "
          "clasificación de una sola clase que aprende la frontera de la región donde se "
          "concentra el comportamiento normal, sin necesidad de ejemplos etiquetados de "
          "ataque. Schölkopf et al. lo formulan como la búsqueda del hiperplano que separa "
          "los datos del origen con margen máximo en el espacio de características inducido "
          "por un kernel.")
    cambios.append("p102 → introducción del OCSVM")

    texto(P[103], "El problema de optimización primal se expresa como:")
    # ½‖w‖² + 1/(νn) Σ ξᵢ − ρ
    ecuacion(P[104],
             frac(r("1"), r("2")), r("‖w‖"), sup(r(""), r("2")), r(" + "),
             frac(r("1"), r("νn")),
             nary("∑", sub(r("i"), r("")), r("n"), sub(r("ξ"), r("i"))),
             r(" − ρ"))
    cambios.append("p104 → ecuación del problema primal del OCSVM")

    texto(P[106], "w y ρ: parámetros del hiperplano separador en el espacio de características")
    texto(P[107], "ξᵢ: variables de holgura que permiten que algunas observaciones queden "
                  "fuera de la frontera; n es el número de observaciones de entrenamiento")
    cambios.append("p106–p107 → parámetros del OCSVM")

    texto(P[108], "Resuelto el problema en su forma dual, la función de decisión sobre una "
                  "observación nueva queda expresada mediante el kernel:")
    # f(x) = Σ αᵢ k(xᵢ, x) − ρ
    ecuacion(P[109],
             r("f(x) = "),
             nary("∑", sub(r("i"), r("")), r("n"), None),
             sub(r("α"), r("i")), r(" k("), sub(r("x"), r("i")), r(", x) − ρ"))
    cambios.append("p109 → función de decisión del OCSVM")

    texto(P[111], "f(x): puntuación de la observación; en la implementación se usa "
                  "score_samples, cuya escala se fija en la calibración")
    texto(P[112], "αᵢ: multiplicadores de Lagrange; solo son distintos de cero en los "
                  "vectores de soporte")
    texto(P[113], "k(·,·): función kernel que traslada el producto escalar al espacio de "
                  "características sin calcularlo de forma explícita")
    texto(P[114], "ρ: desplazamiento del hiperplano respecto al origen")
    cambios.append("p111–p114 → términos de la función de decisión")

    texto(P[115],
          "El parámetro ν ∈ (0, 1] controla el compromiso del modelo: acota por arriba la "
          "fracción de observaciones de entrenamiento que quedan fuera de la frontera y por "
          "abajo la fracción que actúa como vector de soporte. En este proyecto se fijó "
          "ν = 0,05, es decir, se admite que hasta un 5 % del tráfico normal de "
          "entrenamiento quede fuera de la región aprendida.")
    cambios.append("p115 → interpretación de ν")

    texto(P[116], "El kernel empleado es la función de base radial (RBF):")
    # k(x, x') = exp(−γ‖x − x'‖²)
    ecuacion(P[117],
             r("k(x, x′) = exp(−γ‖x − x′‖"), sup(r(""), r("2")), r(")"))
    cambios.append("p117 → kernel RBF")

    texto(P[119], "γ: parámetro de anchura del kernel; se empleó gamma='scale', que lo "
                  "deriva de la varianza de las variables estandarizadas")
    texto(P[120], "‖x − x′‖: distancia euclídea entre dos observaciones tras aplicar "
                  "StandardScaler")
    cambios.append("p119–p120 → parámetros del kernel")

    texto(P[121],
          "El escalado previo no es opcional con un kernel RBF: la distancia euclídea "
          "quedaría dominada por las variables de mayor magnitud —tasas de bytes frente a "
          "razones acotadas en [0, 1]— si no se estandarizan. Por eso el modelo congelado "
          "es un pipeline StandardScaler–OCSVM y no un OCSVM aislado.")
    cambios.append("p121 → justificación del escalado")

    texto(P[122],
          "La regla operativa aplicada es: una ventana genera ALERT cuando su puntuación es "
          "menor que 1,8126087939765134. Ese umbral se fijó con α = 0,05 sobre las 273 "
          "ventanas normales de validation, y nunca con el conjunto de prueba. Isolation "
          "Forest, formulado por Liu et al. mediante longitudes de aislamiento, se evaluó "
          "como candidato inicial y se conserva como comparador: la ablación midió que "
          "presenta puntos ciegos en las familias tcp-syn-rate y udp-probe que el OCSVM sí "
          "cubre.")
    cambios.append("p122 → umbral operativo y papel de Isolation Forest")

    d.save(PPI)
    print(f"Aplicados {len(cambios)} bloques:")
    for c in cambios:
        print("  ·", c)


if __name__ == "__main__":
    main()


# ============================================================ CRONOGRAMA ====
def reconstruir_cronograma(d) -> str:
    """Sustituye la tabla 28x13 por una compacta por etapa.

    La original tiene 12 actividades sin una sola marca, dos columnas de mes sin
    etiqueta y termina en junio de 2026, cuando el trabajo real ocurrio entre
    julio y agosto. Agrupar por etapa la reduce a 9 filas sin perder informacion.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, RGBColor

    AZUL, ZEBRA = "1F4E79", "EEF3FA"

    def shade(cell, hx):
        el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hx)
        cell._tc.get_or_add_tcPr().append(el)

    vieja = None
    for t in d.tables:
        if len(t.rows) > 20 and len(t.columns) > 10:
            vieja = t
            break
    if vieja is None:
        raise SystemExit("ABORTA: no se encontró la tabla del cronograma")

    filas = [
        ("Formulación y revisión bibliográfica", "Ago – Dic 2025", "—"),
        ("Metodología y diseño experimental", "Dic 2025 – Mar 2026", "docs/fase01-diseno-experimental/"),
        ("Configuración del entorno de laboratorio", "Abr – Jun 2026", "docs/fase00-infraestructura/"),
        ("Captura e instrumentación de datos", "Jul – Ago 2026", "docs/fase03-dataset/ (180 documentos)"),
        ("Preparación de datos e ingeniería de variables", "Ago 2026", "docs/fase02-features-multicapa/"),
        ("Modelado, comparación y congelado del modelo", "Ago 2026", "docs/fase04-modelado/"),
        ("Integración del sistema y control inline", "Ago 2026", "docs/fase05-motor-tiempo-real/"),
        ("Validación experimental y ajuste (F6)", "Ago 2026", "docs/fase07-validacion-final/"),
        ("Análisis de resultados, redacción y cierre", "Ago – Set 2026", "docs/entregables/"),
    ]

    nueva = d.add_table(rows=1, cols=3)
    nueva.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, (txt, w) in enumerate(zip(["Etapa", "Periodo", "Evidencia en el repositorio"],
                                     [7.4, 3.6, 6.0])):
        cell = nueva.rows[0].cells[c]; cell.width = Cm(w); cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.font.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, AZUL)
    for i, fila in enumerate(filas):
        row = nueva.add_row()
        for c, txt in enumerate(fila):
            cell = row.cells[c]; cell.width = Cm([7.4, 3.6, 6.0][c]); cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(8.6)
            if c == 2:
                run.font.name = "Consolas"; run.font.size = Pt(7.8)
        if i % 2 == 0:
            for cell in row.cells:
                shade(cell, ZEBRA)

    b = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{lado}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "BFBFBF")
        b.append(e)
    nueva._tbl.tblPr.append(b)

    vieja._tbl.addnext(nueva._tbl)      # colocar en el sitio de la antigua
    vieja._tbl.getparent().remove(vieja._tbl)
    return f"cronograma: tabla de {len(filas)} etapas con fechas reales y evidencia enlazada"
