#!/usr/bin/env python3
"""Justificación de la revista objetivo (Sesión 04, momento CREA).

Formato exigido por la sesión: media página a una página, **aparte de la
matriz** —«Matriz completa + justificación (1 página)», diapositiva 33—.
Las cifras se importan del generador de la matriz para que no puedan
desincronizarse.

    .venv/bin/python3 scripts/entregables/generar_justificacion_revista.py
"""
from __future__ import annotations
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _docx_estilo import rematar                      # noqa: E402
from generar_matriz_revistas import (                 # noqa: E402
    CANDIDATAS, CRITERIOS, ELEGIDAS, H5, VOLUMEN, planes, total, total10)

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "docs" / "entregables" / "assets" / "logo-upeu.png"
OUT = REPO / "docs/entregables/09-matriz-revistas/Justificacion-revista-objetivo.docx"

INK = RGBColor(0x13, 0x1B, 0x2E)
DIM = RGBColor(0x5B, 0x6B, 0x8C)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_HEAD, F_ZEBRA, F_OK = "1F4E79", "EEF3FA", "E0F3E6"


def shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def h1(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(txt)
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    return p


def parrafo(doc, txt, size=9.2, italic=False, color=INK, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i, tramo in enumerate(txt.split("**")):
        r = p.add_run(tramo)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.italic = italic
        r.font.bold = i % 2 == 1
    return p


def tabla(doc, cabeceras, filas, anchos, destacar=0):
    t = doc.add_table(rows=1, cols=len(cabeceras))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c, (txt, w) in enumerate(zip(cabeceras, anchos)):
        cell = t.rows[0].cells[c]
        cell.width = Cm(w)
        cell.text = ""
        r = cell.paragraphs[0].add_run(txt)
        r.font.bold = True
        r.font.size = Pt(8.2)
        r.font.color.rgb = WHITE
        shade(cell, F_HEAD)
    for i, fila in enumerate(filas):
        row = t.add_row()
        for c, txt in enumerate(fila):
            cell = row.cells[c]
            cell.width = Cm(anchos[c])
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            for k, tramo in enumerate(str(txt).split("**")):
                r = p.add_run(tramo)
                r.font.size = Pt(8.0)
                r.font.bold = k % 2 == 1
                r.font.color.rgb = INK
            if i < destacar:
                shade(cell, F_OK)
            elif i % 2 == 1:
                shade(cell, F_ZEBRA)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def main() -> None:
    if not LOGO.exists():
        raise SystemExit(f"falta el logo: {LOGO}")
    P = planes()
    orden = sorted((c for c in CANDIDATAS if c["corto"] in ELEGIDAS),
                   key=lambda c: -total(c))
    pa = orden[0]

    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(1.4)
    s.left_margin = s.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"

    doc.add_picture(str(LOGO), width=Cm(4.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sz, col, bold in [
        ("Universidad Peruana Unión", 10, INK, True),
        ("Facultad de Ingeniería y Arquitectura · E.P. de Ingeniería de Sistemas",
         8.4, DIM, False),
        ("Investigación V · Sesión 04", 8.4, DIM, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.font.color.rgb = col
        r.font.bold = bold

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("JUSTIFICACIÓN DE LA REVISTA OBJETIVO")
    r.font.size = Pt(14.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Sistema open source para la detección temprana de comportamientos "
                  "anómalos en redes de datos\n"
                  "Rubén Mark Salazar Tocas · Uziel Elias Sauñe Fernandez\nAsesores: Ing. Nemias Saboya Ríos · Ing. Fernando Manuel Asin Gómez")
    r.font.size = Pt(8.8)
    r.font.color.rgb = DIM

    # ------------------------------------------------------------- decisión --
    h1(doc, "Decisión")
    parrafo(doc,
            f"Se elige **{pa['nombre']}**. Obtuvo el puntaje ponderado más alto del "
            f"mapa —**{total(pa):.1f}/100**, o **{total10(pa):.2f}/10**— entre nueve "
            f"candidatas legítimas, y encabeza tres de los cinco criterios.")

    filas = []
    for i, c in enumerate(orden):
        h = H5.get(c["corto"])
        filas.append([
            f"**Plan {chr(65+i)}**", c["corto"], f"**{total(c):.1f}**",
            f"{total10(c):.2f}",
            f"{h[0]}" if h else "—",
            f"{VOLUMEN[c['corto']][1]}",
        ])
    tabla(doc, ["", "Revista", "Puntaje /100", "/10", "h5-index", "Art./año"],
          filas, [1.9, 2.2, 2.4, 1.6, 2.2, 2.0], destacar=1)
    parrafo(doc, "El orden de los planes no se eligió: **se calcula desde la matriz**, "
                 "como pide la sesión. La segunda mejor puntuada es el Plan B y así "
                 "sucesivamente.", size=8.4, italic=True, color=DIM)

    # --------------------------------------------------------- por qué esta --
    h1(doc, f"Por qué {pa['corto']} y no las otras")
    tabla(doc, ["Razón", "Evidencia verificada"], [
        ["**Encaje temático**",
         "**86 artículos desde 2024** sobre detección de intrusiones, de anomalías o "
         "seguridad de redes. Ninguna otra se acerca"],
        ["**Coste más bajo**",
         "APC de **USD 300** —USD 400 desde el 1 de octubre de 2026—, por debajo de los "
         "USD 415 de BEEI y de los USD 850 de IJSSE e ISI"],
        ["**El ciclo más corto**",
         "Medido sobre cinco de sus artículos: **41 días** de mediana a la primera "
         "decisión y **158 hasta publicar**. ISI tarda 250"],
        ["**Capacidad real**",
         "**556 artículos en 2025**, mensual. Publica su tasa de aceptación —17,9 %—, "
         "cosa que ninguna otra hace"],
        ["**Legitimidad confirmada**",
         "Ficha de Scopus consultada el 02/09/2026: **«covered by Scopus from 2008 to "
         "2026»**, CiteScore 3,3, Q2 en dos categorías, SJR 0,292. Fuera de la lista de "
         "depredadoras"],
        ["**Coste cubierto**",
         "El **APC lo asume la Universidad Peruana Unión**; enviando antes del 1 de "
         "octubre de 2026 son USD 300 en vez de USD 400"],
    ], [3.3, 13.9])

    # ------------------------------------------------------------- reservas --
    h1(doc, "Lo que juega en contra, dicho por delante")
    tabla(doc, ["Reserva", "Alcance"], [
        ["**Acceso abierto en bronce**",
         "Sus artículos se leen gratis pero **sin licencia declarada** y la revista no "
         "figura en DOAJ. BEEI, el Plan B, es diamante con CC BY-SA: mejor en esto"],
        ["**Sin h5-index**",
         "**No figura en Google Scholar Metrics** —comprobado con tres consultas—. "
         "BEEI tiene 41, ISI 23"],
        ["**Cesión de copyright**",
         "Exige transferir los derechos al editor al aceptar"],
        ["**CiteScore a la baja**",
         "El CiteScoreTracker 2026 va en **2,8** frente al 3,3 de 2025. Hay que "
         "vigilarlo: si cierra por debajo de 3,0 puede perder el Q2 en una categoría"],
    ], [4.2, 13.0])
    parrafo(doc,
            "Ninguna invalida la elección: encaje, coste y plazo pesan el 65 % de la "
            "matriz y en los tres encabeza. Se declaran porque **una decisión defendible "
            "no es la que no tiene puntos débiles, sino la que los tiene medidos**.",
            size=8.8)

    # -------------------------------------------------------------- cascada --
    h1(doc, "Plan de respaldo")
    parrafo(doc,
            f"Ante un rechazo, el manuscrito pasa a **{orden[1]['corto']}** "
            f"({total(orden[1]):.1f}) y luego a **{orden[2]['corto']}** "
            f"({total(orden[2]):.1f}), sin envío simultáneo. La adaptación es de "
            f"formato: la estructura del mapeo de 21 artículos sirve para las cuatro.")
    parrafo(doc,
            "**Verificación abierta:** confirmar por escrito con la coordinación si la "
            "lista de control inhabilita una revista o solo registra lo publicado. De "
            "eso depende que los Planes B y C sigan disponibles.", size=8.8, color=DIM)

    rematar(doc,
            "Justificación de la revista objetivo",
            "Por qué IJIES y no las otras candidatas evaluadas",
            "Justificación de la revista objetivo · Salazar Tocas & Sauñe Fernandez",
            "Investigación V · Sesión 04 · UPeU")
    doc.save(OUT)

    pal = sum(len(p.text.split()) for p in doc.paragraphs)
    pal += sum(len(c.text.split()) for t in doc.tables for r in t.rows for c in r.cells)
    print(f"Generado: {OUT.relative_to(REPO)}")
    print(f"  ~{pal} palabras · ~{pal/450:.1f} páginas (exigido: 0,5 a 1)")


if __name__ == "__main__":
    main()
